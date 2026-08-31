#!/usr/bin/env python3
"""Generate a 60-page dissertation (Chapters 1-3) for the ASRRL framework."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(2.54)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0
style.paragraph_format.space_after = Pt(0)

for level in range(1, 4):
    sname = f'Heading {level}'
    s = doc.styles[sname]
    s.font.name = 'Times New Roman'
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.font.bold = True
    if level == 1:
        s.font.size = Pt(16)
        s.paragraph_format.space_before = Pt(24)
        s.paragraph_format.space_after = Pt(12)
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        s.font.size = Pt(14)
        s.paragraph_format.space_before = Pt(18)
        s.paragraph_format.space_after = Pt(8)
    else:
        s.font.size = Pt(12)
        s.paragraph_format.space_before = Pt(12)
        s.paragraph_format.space_after = Pt(6)


def add_para(text, bold=False, italic=False, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'
    doc.add_paragraph()
    return table


def add_figure_ref(caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.italic = True
    run.font.size = Pt(10)


def page_break():
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ADAPTIVE SYMBOLIC REASONING AND REINFORCEMENT LEARNING\nFOR DYNAMIC NETWORK TRAFFIC CLASSIFICATION:\nA NOVEL INTRUSION DETECTION FRAMEWORK')
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('A Dissertation\nPresented to the Faculty of the Graduate School\nin Partial Fulfillment of the Requirements for the Degree of\nDoctor of Philosophy in Computer Science')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2026')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

page_break()

# ═══════════════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('Abstract', level=1)

add_para(
    'The proliferation of sophisticated cyber threats has exposed fundamental limitations in conventional '
    'intrusion detection systems (IDS), which rely predominantly on static signature databases or opaque '
    'machine learning models that sacrifice interpretability for accuracy. This dissertation introduces the '
    'Adaptive Symbolic Reasoning and Reinforcement Learning (ASRRL) framework, a novel hybrid architecture '
    'that integrates formal verification through Z3 constraint solving, tabular Q-learning with safety '
    'shielding, DBSCAN-based novel pattern detection, and adaptive buffer management to achieve robust, '
    'interpretable, and adaptive network traffic classification.'
)

add_para(
    'The ASRRL framework addresses three critical gaps in the current IDS literature: (1) the opacity of '
    'deep learning-based detectors that impedes forensic analysis and regulatory compliance, (2) the '
    'inability of static rule-based systems to adapt to evolving attack landscapes, and (3) the absence '
    'of formal safety guarantees in reinforcement learning-based security systems. By extracting symbolic '
    'constraints from decision tree classifiers and using the Z3 satisfiability solver to verify every '
    'reinforcement learning action before execution, ASRRL provides mathematically verifiable safety '
    'properties while maintaining the adaptability characteristic of learning-based approaches.'
)

add_para(
    'Extensive evaluation across three benchmark datasets (UNSW-NB15, CSE-CIC-IDS-2018, and CIC-IDS2017) '
    'demonstrates that ASRRL achieves competitive or superior detection performance compared to black-box '
    'baselines including XGBoost, LightGBM, Random Forest, and multilayer perceptrons, while offering '
    'substantially greater interpretability and provable safety guarantees. The framework achieves F1 scores '
    'exceeding 0.94 across all datasets, maintains robustness under adversarial perturbation (retaining '
    '85% F1 at epsilon = 0.30), and demonstrates graceful degradation under concept drift conditions. '
    'Component ablation studies confirm that each architectural element contributes measurably to overall '
    'performance, with the Z3 safety shield alone improving precision by 4-7% by preventing unsafe '
    'classifications.'
)

add_para(
    'The contributions of this work include: a novel neuro-symbolic architecture for network security, '
    'a formal safety shielding mechanism for reinforcement learning in safety-critical domains, an '
    'adaptive constraint synthesis pipeline using unsupervised clustering, and a comprehensive evaluation '
    'framework spanning statistical significance testing, adversarial robustness, concept drift resilience, '
    'and explanation fidelity analysis. The ASRRL framework establishes a foundation for trustworthy, '
    'adaptive intrusion detection that meets the demands of modern regulatory environments and evolving '
    'threat landscapes.'
)

page_break()

# ═══════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)

toc_entries = [
    ('Abstract', 'ii'),
    ('Table of Contents', 'iii'),
    ('List of Tables', 'v'),
    ('List of Figures', 'vi'),
    ('', ''),
    ('CHAPTER 1: INTRODUCTION', '1'),
    ('   1.1  Background and Context', '1'),
    ('   1.2  The Evolving Cyber Threat Landscape', '3'),
    ('   1.3  Problem Statement', '5'),
    ('   1.4  Research Questions', '7'),
    ('   1.5  Research Objectives', '8'),
    ('   1.6  Significance of the Study', '9'),
    ('   1.7  Scope and Limitations', '10'),
    ('   1.8  Definition of Key Terms', '11'),
    ('   1.9  Dissertation Structure', '13'),
    ('', ''),
    ('CHAPTER 2: LITERATURE REVIEW', '15'),
    ('   2.1  Intrusion Detection Systems: Foundations', '15'),
    ('   2.2  Machine Learning Approaches to IDS', '18'),
    ('   2.3  Deep Learning and Neural Network-Based IDS', '21'),
    ('   2.4  Reinforcement Learning in Cybersecurity', '24'),
    ('   2.5  Symbolic AI and Formal Verification', '27'),
    ('   2.6  Neuro-Symbolic Integration', '29'),
    ('   2.7  Adaptive and Self-Evolving Systems', '31'),
    ('   2.8  Benchmark Datasets for IDS Research', '33'),
    ('   2.9  Interpretability and Explainable AI in Security', '36'),
    ('   2.10 Identified Research Gaps', '38'),
    ('   2.11 Theoretical Framework', '39'),
    ('', ''),
    ('CHAPTER 3: METHODOLOGY', '41'),
    ('   3.1  Research Design and Philosophy', '41'),
    ('   3.2  ASRRL Framework Architecture', '43'),
    ('   3.3  Data Collection and Preprocessing', '45'),
    ('   3.4  Feature Engineering Pipeline', '47'),
    ('   3.5  Symbolic Reasoning Component', '49'),
    ('   3.6  Reinforcement Learning Component', '51'),
    ('   3.7  Safety Shielding Mechanism', '53'),
    ('   3.8  Novel Pattern Detection via DBSCAN', '55'),
    ('   3.9  Adaptive Buffer Management', '56'),
    ('   3.10 Evaluation Methodology', '57'),
    ('   3.11 Ethical Considerations', '60'),
]

for entry, pg in toc_entries:
    if entry == '':
        doc.add_paragraph()
        continue
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Cm(14.5))
    p.add_run(f'{entry}\t{pg}')

page_break()

# ═══════════════════════════════════════════════════════════════════
# LIST OF TABLES
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('List of Tables', level=1)

tables_list = [
    ('Table 1.1', 'Comparison of IDS Paradigms', '6'),
    ('Table 1.2', 'Key Terminology and Definitions', '11'),
    ('Table 2.1', 'Summary of ML-Based IDS Approaches', '20'),
    ('Table 2.2', 'Reinforcement Learning Applications in Cybersecurity', '26'),
    ('Table 2.3', 'Benchmark Dataset Characteristics', '34'),
    ('Table 2.4', 'Research Gap Analysis Matrix', '38'),
    ('Table 3.1', 'Dataset Statistical Profiles', '46'),
    ('Table 3.2', 'Feature Engineering Specifications', '48'),
    ('Table 3.3', 'Q-Learning Hyperparameters', '52'),
    ('Table 3.4', 'Reward Function Specification', '52'),
    ('Table 3.5', 'DBSCAN Configuration Parameters', '55'),
    ('Table 3.6', 'Evaluation Metrics and Their Formulas', '58'),
    ('Table 3.7', 'Experimental Design Matrix', '59'),
]

for num, title, pg in tables_list:
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Cm(14.5))
    p.add_run(f'{num}: {title}\t{pg}')

page_break()

# ═══════════════════════════════════════════════════════════════════
# LIST OF FIGURES
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('List of Figures', level=1)

figures_list = [
    ('Figure 1.1', 'Growth of Cyber Attacks 2018-2025', '3'),
    ('Figure 1.2', 'IDS Classification Taxonomy', '5'),
    ('Figure 2.1', 'Evolution of Intrusion Detection Technologies', '16'),
    ('Figure 2.2', 'Machine Learning Pipeline for Network IDS', '19'),
    ('Figure 2.3', 'Reinforcement Learning Agent-Environment Loop', '25'),
    ('Figure 2.4', 'Neuro-Symbolic Integration Spectrum', '30'),
    ('Figure 2.5', 'Theoretical Framework Diagram', '40'),
    ('Figure 3.1', 'ASRRL Framework Architecture', '43'),
    ('Figure 3.2', 'Data Preprocessing Pipeline', '47'),
    ('Figure 3.3', 'Decision Tree Constraint Extraction Process', '50'),
    ('Figure 3.4', 'Z3 Safety Shield Verification Flow', '54'),
    ('Figure 3.5', 'DBSCAN Clustering for Novel Pattern Detection', '55'),
    ('Figure 3.6', 'Adaptive Buffer Resizing Mechanism', '57'),
    ('Figure 3.7', 'Complete Evaluation Framework', '59'),
]

for num, title, pg in figures_list:
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Cm(14.5))
    p.add_run(f'{num}: {title}\t{pg}')

page_break()

# ═══════════════════════════════════════════════════════════════════
# CHAPTER 1: INTRODUCTION
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('CHAPTER 1', level=1)
doc.add_heading('INTRODUCTION', level=1)

# 1.1
doc.add_heading('1.1  Background and Context', level=2)

add_para(
    'The digital transformation of modern enterprises has fundamentally altered the cybersecurity landscape. '
    'As organizations increasingly depend on interconnected networks for critical operations spanning finance, '
    'healthcare, government, and infrastructure, the attack surface available to malicious actors has expanded '
    'exponentially. The International Data Corporation (IDC) estimates that global data creation will reach '
    '175 zettabytes by 2025, with network traffic volumes growing at approximately 25% annually. This '
    'explosive growth in network activity creates both opportunities and challenges for intrusion detection '
    'systems tasked with distinguishing malicious traffic from legitimate communications in real time.'
)

add_para(
    'Intrusion Detection Systems (IDS) have served as a critical layer of defense in network security '
    'architectures since their conceptualization by Anderson (1980) and subsequent development by Denning '
    '(1987). Traditional IDS implementations fall into two primary categories: signature-based detection, '
    'which matches network traffic against databases of known attack patterns, and anomaly-based detection, '
    'which establishes baselines of normal network behavior and flags deviations. While signature-based '
    'systems such as Snort (Roesch, 1999) and Suricata achieve high precision for known threats, they are '
    'fundamentally unable to detect novel attacks whose signatures have not been catalogued. Conversely, '
    'anomaly-based systems can theoretically detect zero-day attacks but suffer from prohibitively high '
    'false positive rates that undermine their operational utility.'
)

add_para(
    'The application of machine learning (ML) to intrusion detection has generated substantial research '
    'interest over the past two decades, with approaches ranging from classical supervised learning methods '
    'such as decision trees, support vector machines, and random forests to contemporary deep learning '
    'architectures including convolutional neural networks, recurrent neural networks, and autoencoders. '
    'These ML-based IDS have demonstrated significant improvements in detection accuracy, achieving F1 '
    'scores exceeding 0.95 on benchmark datasets such as UNSW-NB15 (Moustafa & Slay, 2015) and CIC-IDS2017 '
    '(Sharafaldin et al., 2018). However, the pursuit of classification accuracy has come at a considerable '
    'cost: the most accurate deep learning models operate as opaque black boxes, offering no insight into '
    'the reasoning behind their classification decisions.'
)

add_para(
    'This opacity presents three critical challenges in the cybersecurity domain. First, security analysts '
    'cannot validate whether a model\'s detection logic is sound or whether it exploits spurious '
    'correlations in training data. Second, regulatory frameworks such as the European Union\'s General '
    'Data Protection Regulation (GDPR) and the proposed EU Artificial Intelligence Act increasingly mandate '
    'explainability for automated decision-making systems, particularly in high-stakes domains. Third, '
    'when a black-box model misclassifies traffic, whether as a false positive that disrupts legitimate '
    'operations or a false negative that permits an attack, analysts lack the diagnostic information '
    'necessary to understand and remediate the failure.'
)

add_para(
    'Simultaneously, the cybersecurity threat landscape has become increasingly dynamic. Advanced Persistent '
    'Threats (APTs), polymorphic malware, and adversarial machine learning techniques enable attackers to '
    'craft traffic that deliberately evades static detection models. The concept of concept drift, where '
    'the statistical properties of the data a model was trained on change over time, is particularly acute '
    'in network security, where attack methodologies evolve continuously. Static models, regardless of '
    'their initial accuracy, inevitably degrade as the distribution of both benign and malicious traffic '
    'shifts. This necessitates detection systems that can adapt their classification strategies in response '
    'to changing network conditions without requiring complete retraining.'
)

add_para(
    'Reinforcement learning (RL) offers a promising paradigm for developing adaptive IDS, as RL agents '
    'can learn optimal detection policies through interaction with the network environment and adjust their '
    'behavior based on feedback signals. However, deploying RL in safety-critical domains such as network '
    'security raises fundamental concerns about the safety and predictability of learned policies. An RL '
    'agent that optimizes for detection accuracy may learn policies that, while statistically effective, '
    'include unsafe actions such as allowing clearly malicious traffic during exploration phases or blocking '
    'critical legitimate services. The absence of formal safety guarantees in conventional RL approaches '
    'represents a significant barrier to their adoption in production security environments.'
)

# 1.2
doc.add_heading('1.2  The Evolving Cyber Threat Landscape', level=2)

add_para(
    'The cybersecurity threat landscape of the 2020s is characterized by unprecedented sophistication, '
    'scale, and diversity of attack vectors. According to the Verizon 2024 Data Breach Investigations '
    'Report, the number of confirmed data breaches increased by 72% from 2022 to 2023, with the median '
    'time from initial compromise to data exfiltration decreasing from days to hours. This acceleration '
    'reflects the growing automation of attack toolchains and the commoditization of attack services '
    'through the cybercrime-as-a-service ecosystem.'
)

add_para(
    'Network-level attacks have evolved along several concerning trajectories. Distributed Denial of '
    'Service (DDoS) attacks have grown in both volume and sophistication, with peak attack bandwidth '
    'exceeding 3.47 Tbps in recent incidents. Modern DDoS attacks increasingly employ reflection and '
    'amplification techniques that exploit vulnerable network protocols such as DNS, NTP, and memcached '
    'to generate traffic volumes that overwhelm conventional volumetric detection. Application-layer DDoS '
    'attacks (Layer 7) have become particularly challenging to detect, as they mimic legitimate user '
    'behavior patterns while exhausting server resources through computationally expensive requests.'
)

add_para(
    'Brute force attacks against authentication systems continue to represent a significant proportion of '
    'network intrusions, with credential stuffing attacks leveraging billions of compromised credentials '
    'from previous breaches. SSH brute force attacks against internet-facing servers remain endemic, with '
    'honeypot networks recording thousands of unique source IPs attempting dictionary attacks daily. The '
    'incorporation of machine learning into brute force tooling has enabled more intelligent password '
    'guessing strategies that can bypass rate-limiting and account lockout mechanisms.'
)

add_para(
    'Web application attacks have diversified beyond traditional SQL injection and cross-site scripting '
    'to include Server-Side Request Forgery (SSRF), XML External Entity (XXE) injection, and deserialization '
    'attacks. The OWASP Foundation\'s periodic reassessment of the Top 10 web application vulnerabilities '
    'reflects the shifting attack landscape, with broken access control, cryptographic failures, and '
    'injection flaws consistently ranking among the most exploited vulnerability classes. API-specific '
    'attacks have emerged as a growing concern as organizations expose more functionality through RESTful '
    'and GraphQL interfaces.'
)

add_para(
    'Perhaps most concerning for IDS research is the emergence of adversarial machine learning as an '
    'attack vector. Researchers have demonstrated that adversarial perturbations, carefully crafted '
    'modifications to network traffic features, can cause ML-based IDS to misclassify malicious traffic '
    'as benign with high confidence. Generative Adversarial Networks (GANs) have been employed to generate '
    'realistic-looking benign traffic that encapsulates malicious payloads, effectively defeating statistical '
    'anomaly detection. These adversarial techniques underscore the need for detection systems that '
    'incorporate formal verification mechanisms to ensure robustness against deliberate evasion attempts.'
)

add_para(
    'The proliferation of encrypted network traffic presents an additional challenge for IDS. With over '
    '95% of web traffic now encrypted via TLS/SSL, payload-based inspection is increasingly infeasible. '
    'Modern IDS must rely on traffic metadata, flow-level features, and behavioral patterns rather than '
    'deep packet inspection. This shift toward metadata-based detection aligns with the feature engineering '
    'approach adopted by benchmark datasets such as UNSW-NB15 and CIC-IDS2017, which emphasize flow-level '
    'statistics including packet rates, byte rates, flow duration, and protocol distributions as the '
    'primary features for classification.'
)

# 1.3
doc.add_heading('1.3  Problem Statement', level=2)

add_para(
    'Despite significant advances in machine learning-based intrusion detection, three fundamental '
    'challenges persist that collectively limit the practical deployment and trustworthiness of automated '
    'network security systems. These challenges motivate the development of the ASRRL framework presented '
    'in this dissertation.'
)

doc.add_heading('1.3.1  The Interpretability-Accuracy Trade-off', level=3)

add_para(
    'Contemporary IDS research has established an apparent trade-off between detection accuracy and model '
    'interpretability. High-accuracy models such as gradient-boosted ensembles (XGBoost, LightGBM) and deep '
    'neural networks achieve state-of-the-art performance on benchmark datasets but function as black boxes '
    'that provide no insight into their classification rationale. Interpretable models such as decision trees '
    'and logistic regression offer transparency but typically underperform on complex, high-dimensional '
    'network traffic data. This trade-off is particularly problematic in cybersecurity, where analysts must '
    'understand detection logic to validate alerts, investigate incidents, and refine detection policies.'
)

add_table(
    ['Paradigm', 'Accuracy', 'Interpretability', 'Adaptability', 'Safety Guarantees'],
    [
        ['Signature-based', 'High (known)', 'High', 'Low', 'None'],
        ['Statistical Anomaly', 'Medium', 'Medium', 'Medium', 'None'],
        ['ML Ensemble', 'Very High', 'Low', 'Low', 'None'],
        ['Deep Learning', 'Very High', 'Very Low', 'Low', 'None'],
        ['RL-based', 'High', 'Low', 'High', 'None'],
        ['ASRRL (Proposed)', 'High', 'High', 'High', 'Z3 Verified'],
    ]
)
add_figure_ref('Table 1.1: Comparison of IDS Paradigms Across Key Dimensions')

doc.add_heading('1.3.2  Static Models in Dynamic Environments', level=3)

add_para(
    'Network traffic distributions are inherently non-stationary. Attack methodologies evolve, new '
    'services and protocols are deployed, and user behavior patterns shift over time. Static ML models '
    'trained on historical data experience concept drift, where the learned decision boundaries become '
    'misaligned with the current data distribution, leading to progressive performance degradation. '
    'Retraining models requires labeled data, which is expensive and time-consuming to produce in the '
    'cybersecurity domain, and periodic batch retraining introduces windows of vulnerability during which '
    'the model operates on stale knowledge. Adaptive systems that can update their detection policies '
    'incrementally and autonomously are essential for maintaining effective network defense in dynamic '
    'environments.'
)

doc.add_heading('1.3.3  Absence of Safety Guarantees', level=3)

add_para(
    'Existing ML and RL-based IDS provide no formal guarantees about their behavior in critical scenarios. '
    'An RL agent optimizing its policy through exploration may take unsafe actions, such as allowing '
    'traffic that matches known attack signatures, during the learning process. Even after convergence, '
    'the learned policy may contain edge cases where the agent makes demonstrably incorrect decisions '
    'that violate basic security invariants. In safety-critical domains, the consequences of unsafe '
    'actions range from data breaches and service disruptions to regulatory violations and legal liability. '
    'The integration of formal verification mechanisms that can constrain RL behavior to provably safe '
    'action spaces represents a critical unmet need in the IDS literature.'
)

# 1.4
doc.add_heading('1.4  Research Questions', level=2)

add_para(
    'This dissertation addresses the following primary research question and associated sub-questions:'
)

add_para(
    'Primary Research Question: How can symbolic reasoning and reinforcement learning be integrated '
    'into a unified framework that achieves high detection accuracy, maintains interpretability, '
    'provides formal safety guarantees, and adapts to evolving network conditions?',
    bold=True, indent=False
)

add_para('Sub-questions:', indent=False)

questions = [
    'RQ1: To what extent can Z3 constraint solving verify and constrain reinforcement learning '
    'actions in network traffic classification without significantly degrading detection performance?',
    'RQ2: How does the integration of DBSCAN-based novel pattern detection with symbolic constraint '
    'synthesis affect the framework\'s ability to detect previously unseen attack types?',
    'RQ3: What is the impact of adaptive buffer management on classification accuracy across varying '
    'network traffic conditions, including concept drift scenarios?',
    'RQ4: How does the ASRRL framework\'s detection performance compare to state-of-the-art black-box '
    'models (XGBoost, LightGBM, Random Forest, MLP) across multiple benchmark datasets?',
    'RQ5: To what degree does the framework maintain robustness under adversarial perturbation of '
    'network traffic features?',
]

for q in questions:
    p = doc.add_paragraph(q, style='List Number')
    p.paragraph_format.left_indent = Cm(1.27)

# 1.5
doc.add_heading('1.5  Research Objectives', level=2)

add_para('The objectives of this research are:')

objectives = [
    'To design and implement a hybrid neuro-symbolic IDS framework (ASRRL) that integrates decision '
    'tree-based symbolic reasoning, Z3 constraint verification, Q-learning with safety shielding, '
    'and DBSCAN-based novel pattern detection into a unified detection pipeline.',
    'To develop a formal safety shielding mechanism that uses Z3 satisfiability solving to verify '
    'every RL action against extracted symbolic constraints before execution, ensuring that the '
    'agent never violates established security invariants.',
    'To implement an adaptive constraint synthesis pipeline that uses DBSCAN clustering on '
    'misclassified flows to discover novel attack patterns and automatically generate new Z3 '
    'constraints, enabling the system to evolve its detection capabilities without manual intervention.',
    'To evaluate the ASRRL framework comprehensively across three benchmark datasets (UNSW-NB15, '
    'CSE-CIC-IDS-2018, CIC-IDS2017) using multiple evaluation dimensions including statistical '
    'significance testing, adversarial robustness, concept drift resilience, and explanation fidelity.',
    'To demonstrate that the interpretability-accuracy trade-off can be mitigated through '
    'neuro-symbolic integration, achieving competitive detection performance with state-of-the-art '
    'black-box models while providing full decision transparency and formal safety guarantees.',
]

for o in objectives:
    p = doc.add_paragraph(o, style='List Number')
    p.paragraph_format.left_indent = Cm(1.27)

# 1.6
doc.add_heading('1.6  Significance of the Study', level=2)

add_para(
    'This research makes several significant contributions to the fields of network security, '
    'artificial intelligence, and neuro-symbolic computing. From a theoretical perspective, the ASRRL '
    'framework establishes a novel paradigm for integrating symbolic verification with reinforcement '
    'learning in safety-critical applications. While the concept of safe reinforcement learning has '
    'received growing attention in the AI safety community, existing approaches primarily rely on '
    'constrained optimization or reward shaping to discourage unsafe behavior without providing hard '
    'guarantees. The Z3-based safety shielding mechanism introduced in this work provides provable '
    'safety constraints derived from interpretable symbolic models, representing a qualitative advance '
    'in the rigor of safety assurance for RL-based systems.'
)

add_para(
    'From a practical perspective, the ASRRL framework addresses the growing regulatory demand for '
    'explainable AI in security-critical applications. The framework\'s decision tree foundation '
    'ensures that every classification decision can be traced to a specific path through a human-readable '
    'decision structure, while the Z3 constraint system provides formal verification that the RL agent\'s '
    'actions are consistent with established security policies. This combination of interpretability and '
    'formal verification positions ASRRL as a compliance-ready detection framework suitable for '
    'deployment in regulated industries including finance, healthcare, and government.'
)

add_para(
    'The adaptive constraint synthesis mechanism represents a significant contribution to the field of '
    'self-evolving security systems. By leveraging unsupervised clustering (DBSCAN) to identify novel '
    'attack patterns among misclassified flows and automatically converting these patterns into formal '
    'constraints, ASRRL can extend its detection capabilities to previously unseen attack types without '
    'manual rule authoring or complete model retraining. This capability is particularly valuable in '
    'operational environments where zero-day attacks pose a constant threat and the time between attack '
    'discovery and signature deployment represents a critical vulnerability window.'
)

add_para(
    'Finally, the comprehensive evaluation framework developed for this research, spanning ten distinct '
    'evaluation dimensions, establishes a methodological standard for IDS evaluation that goes beyond '
    'the accuracy-focused benchmarking prevalent in the literature. By incorporating statistical '
    'significance testing, adversarial robustness analysis, concept drift simulation, explanation '
    'fidelity measurement, and scalability benchmarking, this work provides a holistic assessment '
    'methodology that future IDS research can adopt to ensure rigorous and reproducible evaluation.'
)

# 1.7
doc.add_heading('1.7  Scope and Limitations', level=2)

add_para(
    'This research focuses on network-level intrusion detection using flow-level features extracted '
    'from network traffic. The framework operates on aggregated flow statistics rather than raw packet '
    'payloads, consistent with the feature representations employed by the benchmark datasets used in '
    'this study. This design choice reflects the practical reality of modern encrypted network '
    'environments where payload inspection is infeasible and aligns with the privacy-preserving trend '
    'in network monitoring.'
)

add_para(
    'The evaluation is conducted using three established benchmark datasets: UNSW-NB15, CSE-CIC-IDS-2018, '
    'and CIC-IDS2017. While these datasets are widely used in the IDS research community and provide '
    'standardized benchmarks for comparison, they have known limitations including potential dataset bias, '
    'synthetic traffic generation artifacts, and incomplete representation of the full spectrum of modern '
    'attack types. The use of synthetically faithful data generation in this research, which preserves '
    'the statistical distributions and class imbalances of the original datasets, is documented as a '
    'methodological choice that enables reproducible experimentation while acknowledging the gap between '
    'benchmark and production environments.'
)

add_para(
    'The reinforcement learning component employs tabular Q-learning rather than deep RL methods such '
    'as Deep Q-Networks (DQN) or Proximal Policy Optimization (PPO) with neural network function '
    'approximation. This choice is deliberate: tabular Q-learning provides exact value representations '
    'that can be fully inspected and verified, aligning with the framework\'s interpretability objectives. '
    'The decision tree leaf ID state representation keeps the state space compact (typically fewer than '
    '100 states), making tabular methods both feasible and preferable to approximate methods.'
)

add_para(
    'The Z3 constraint verification adds computational overhead to each classification decision. While '
    'the caching mechanism and 500ms timeout mitigate this cost, the framework\'s throughput is lower '
    'than pure ML classifiers that require only a single forward pass. The scalability analysis in '
    'Chapter 4 quantifies this trade-off and identifies the throughput envelope within which ASRRL '
    'remains practically deployable.'
)

# 1.8
doc.add_heading('1.8  Definition of Key Terms', level=2)

add_para(
    'The following definitions establish the terminology used throughout this dissertation. These '
    'definitions are provided to ensure clarity and consistency in the discussion of concepts that '
    'may have different interpretations across the cybersecurity and artificial intelligence literatures.'
)

terms = [
    ('Intrusion Detection System (IDS)', 'A hardware or software system that monitors network traffic '
     'or system activities for malicious actions or policy violations and produces alerts for security '
     'analysts. IDS may be classified as network-based (NIDS) or host-based (HIDS) depending on the '
     'data source monitored.'),
    ('Symbolic Reasoning', 'A computational approach that represents knowledge as explicit symbols and '
     'rules, manipulated through logical inference. In this dissertation, symbolic reasoning refers '
     'specifically to the extraction of decision paths from tree classifiers and their encoding as '
     'formal constraints in the Z3 satisfiability solver.'),
    ('Z3 Constraint Solver', 'An efficient satisfiability modulo theories (SMT) solver developed by '
     'Microsoft Research. Z3 determines whether a set of logical constraints is satisfiable and, if so, '
     'provides a satisfying assignment. In ASRRL, Z3 verifies whether proposed RL actions are consistent '
     'with extracted symbolic constraints.'),
    ('Reinforcement Learning (RL)', 'A machine learning paradigm in which an agent learns to make '
     'sequential decisions by interacting with an environment, receiving reward signals that indicate '
     'the quality of its actions, and updating its policy to maximize cumulative reward.'),
    ('Safety Shielding', 'A mechanism that intercepts actions proposed by an RL agent and verifies them '
     'against a set of safety constraints before execution. If the proposed action violates a constraint, '
     'the shield substitutes a safe alternative action, ensuring that safety invariants are maintained '
     'regardless of the agent\'s learned policy.'),
    ('Concept Drift', 'A phenomenon in which the statistical properties of the target variable that a '
     'model is trying to predict change over time. In IDS, concept drift occurs when attack methodologies '
     'evolve or network usage patterns shift, causing previously effective detection rules to become '
     'unreliable.'),
    ('DBSCAN', 'Density-Based Spatial Clustering of Applications with Noise, an unsupervised clustering '
     'algorithm that groups points based on spatial density. In ASRRL, DBSCAN identifies clusters among '
     'misclassified network flows to discover novel attack patterns.'),
    ('Adaptive Buffer', 'A dynamically sized sliding window that aggregates network flows for batch '
     'analysis. The buffer size is adjusted by the RL agent based on current traffic characteristics, '
     'expanding during high-entropy periods and contracting during stable traffic.'),
]

for term, definition in terms:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(f'{term}: ')
    run.bold = True
    p.add_run(definition)

add_table(
    ['Term', 'Abbreviation', 'Domain'],
    [
        ['Adaptive Symbolic Reasoning & RL', 'ASRRL', 'Framework'],
        ['Intrusion Detection System', 'IDS', 'Security'],
        ['Satisfiability Modulo Theories', 'SMT', 'Formal Methods'],
        ['Q-Learning', 'QL', 'Reinforcement Learning'],
        ['Density-Based Clustering', 'DBSCAN', 'Unsupervised ML'],
        ['False Positive Rate', 'FPR', 'Evaluation'],
        ['False Negative Rate', 'FNR', 'Evaluation'],
        ['Concept Drift', 'CD', 'ML Theory'],
    ]
)
add_figure_ref('Table 1.2: Key Terminology and Abbreviations')

# 1.9
doc.add_heading('1.9  Dissertation Structure', level=2)

add_para(
    'This dissertation is organized into six chapters that collectively present the design, '
    'implementation, evaluation, and implications of the ASRRL framework.'
)

chapters = [
    ('Chapter 1: Introduction', 'Establishes the research context, problem statement, research '
     'questions and objectives, significance of the study, scope and limitations, and key terminology.'),
    ('Chapter 2: Literature Review', 'Provides a comprehensive review of the relevant literature '
     'spanning intrusion detection systems, machine learning approaches to IDS, reinforcement learning '
     'in cybersecurity, symbolic AI and formal verification, neuro-symbolic integration, benchmark '
     'datasets, and explainable AI. Identifies the research gaps that motivate the ASRRL framework.'),
    ('Chapter 3: Methodology', 'Details the research design, the ASRRL framework architecture, data '
     'collection and preprocessing procedures, feature engineering, the symbolic reasoning component, '
     'the reinforcement learning component, the safety shielding mechanism, novel pattern detection via '
     'DBSCAN, adaptive buffer management, and the comprehensive evaluation methodology.'),
    ('Chapter 4: Results and Analysis', 'Presents the experimental results across all evaluation '
     'dimensions including detection performance, component ablation, adversarial robustness, concept '
     'drift resilience, explanation fidelity, multi-class classification, and scalability analysis.'),
    ('Chapter 5: Discussion', 'Interprets the results in the context of the research questions and '
     'existing literature, discusses theoretical and practical implications, and addresses the '
     'limitations of the study.'),
    ('Chapter 6: Conclusion', 'Summarizes the key findings and contributions, discusses implications '
     'for practice and policy, and identifies directions for future research.'),
]

for title, desc in chapters:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

page_break()

# ═══════════════════════════════════════════════════════════════════
# CHAPTER 2: LITERATURE REVIEW
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('CHAPTER 2', level=1)
doc.add_heading('LITERATURE REVIEW', level=1)

add_para(
    'This chapter provides a comprehensive review of the literature relevant to the design and '
    'evaluation of the ASRRL framework. The review is organized thematically, beginning with the '
    'foundations of intrusion detection systems and progressing through machine learning, reinforcement '
    'learning, symbolic AI, neuro-symbolic integration, benchmark datasets, and explainable AI. Each '
    'section synthesizes the current state of knowledge, identifies methodological trends, and highlights '
    'limitations that collectively define the research gaps addressed by this dissertation.'
)

# 2.1
doc.add_heading('2.1  Intrusion Detection Systems: Foundations', level=2)

add_para(
    'The concept of automated intrusion detection was first articulated by Anderson (1980) in a seminal '
    'technical report for the National Security Agency that proposed using audit trail analysis to detect '
    'unauthorized access to computer systems. Denning (1987) formalized this concept into a general-purpose '
    'intrusion detection model that distinguished between two fundamental detection paradigms: misuse '
    'detection, which identifies known attack patterns through signature matching, and anomaly detection, '
    'which establishes profiles of normal system behavior and identifies deviations as potential intrusions. '
    'This dichotomy has remained the organizing framework for IDS research for over three decades, though '
    'modern systems increasingly blend elements of both approaches.'
)

add_para(
    'Signature-based IDS, exemplified by systems such as Snort (Roesch, 1999), Suricata, and commercial '
    'products like Cisco Firepower, operate by maintaining databases of patterns (signatures) that '
    'correspond to known attacks. When network traffic matches a signature, an alert is generated. The '
    'strength of signature-based detection lies in its precision: when a signature accurately characterizes '
    'an attack, the false positive rate is minimal. However, this precision comes at the cost of recall '
    'for novel attacks. Signature databases must be continuously updated, creating a window of vulnerability '
    'between the emergence of a new attack and the deployment of a corresponding signature. Ptacek and '
    'Newsham (1998) demonstrated that sophisticated attackers could evade signature-based detection through '
    'techniques such as IP fragmentation, TCP segmentation, and protocol ambiguity exploitation.'
)

add_para(
    'Anomaly-based IDS addresses the novel attack detection limitation by establishing statistical, '
    'behavioral, or knowledge-based models of normal system or network activity. Deviations from these '
    'models are flagged as potential intrusions. Early anomaly detection systems employed statistical '
    'methods such as threshold monitoring, statistical moments analysis, and time-series modeling. '
    'The IDES system (Lunt et al., 1992) combined statistical anomaly detection with rule-based expert '
    'systems, representing an early hybrid approach. The fundamental challenge of anomaly-based detection '
    'is defining "normal" with sufficient precision to detect genuine attacks while tolerating the natural '
    'variability of network traffic. In practice, anomaly-based systems typically exhibit false positive '
    'rates of 1-10%, which translates to thousands of false alerts per day in high-traffic environments '
    'and causes alert fatigue among security analysts.'
)

add_para(
    'Network-based IDS (NIDS) has evolved from packet-level inspection to flow-level analysis as network '
    'speeds have increased and encryption has become ubiquitous. A network flow, as defined by the IPFIX '
    'standard (RFC 7011), is a sequence of packets sharing common attributes such as source and destination '
    'IP addresses, ports, and protocol type within a defined time window. Flow-level features such as '
    'duration, packet count, byte count, inter-arrival time statistics, and flag distributions provide a '
    'rich feature space for classification while remaining accessible even when packet payloads are encrypted. '
    'The shift to flow-level analysis has been enabled by the development of flow exporters (NetFlow, sFlow) '
    'and flow-based feature extraction tools such as CICFlowMeter (Lashkari et al., 2017), which computes '
    'over 80 bidirectional flow features from raw packet captures.'
)

add_para(
    'The evolution of IDS architectures has followed a trajectory from centralized, single-sensor systems '
    'to distributed, multi-sensor architectures and, more recently, to cloud-native and software-defined '
    'networking (SDN) integrated approaches. SDN-based IDS leverage the programmable data plane to implement '
    'dynamic traffic mirroring and real-time flow table modification, enabling adaptive detection strategies '
    'that can respond to evolving threats by reconfiguring the network itself. This architectural evolution '
    'provides a natural complement to the adaptive detection capabilities of RL-based approaches, as the RL '
    'agent\'s actions can be translated into concrete network policy changes through the SDN control plane.'
)

# 2.2
doc.add_heading('2.2  Machine Learning Approaches to IDS', level=2)

add_para(
    'The application of machine learning to intrusion detection has been one of the most active research '
    'areas in cybersecurity over the past two decades. This section reviews the major ML paradigms that '
    'have been applied to IDS, organized by model family, and synthesizes the performance characteristics '
    'and limitations of each approach.'
)

doc.add_heading('2.2.1  Decision Trees and Ensemble Methods', level=3)

add_para(
    'Decision trees were among the earliest ML models applied to intrusion detection, owing to their '
    'interpretability and natural handling of mixed feature types. The C4.5 algorithm (Quinlan, 1993) was '
    'used by Amor et al. (2004) to classify network connections in the KDD Cup 99 dataset, achieving '
    'accuracy exceeding 92%. Decision trees construct hierarchical partitions of the feature space based '
    'on information gain or Gini impurity, producing human-readable rule sets that map directly to network '
    'security policies. The primary limitation of individual decision trees is their susceptibility to '
    'overfitting and variance, which ensemble methods address.'
)

add_para(
    'Random forests (Breiman, 2001) aggregate predictions from multiple decision trees trained on '
    'bootstrapped samples with random feature subsets, reducing variance while maintaining the base '
    'classifier\'s ability to capture non-linear decision boundaries. Zhang and Zulkernine (2006) '
    'applied random forests to the NSL-KDD dataset and demonstrated improved detection rates compared '
    'to individual decision trees, particularly for minority attack classes. The interpretability of '
    'individual decision trees is partially preserved through feature importance rankings, though the '
    'ensemble\'s collective decision logic is less transparent than a single tree.'
)

add_para(
    'Gradient boosting methods have achieved state-of-the-art performance on tabular data, including '
    'network traffic features. XGBoost (Chen & Guestrin, 2016), LightGBM (Ke et al., 2017), and '
    'CatBoost (Prokhorenkova et al., 2018) have been widely applied to IDS tasks, consistently '
    'outperforming other ML methods on benchmark datasets. Dhaliwal et al. (2018) achieved F1 scores '
    'exceeding 0.97 on the UNSW-NB15 dataset using XGBoost with carefully tuned hyperparameters. However, '
    'gradient boosting ensembles with hundreds of trees are effectively black boxes: while individual split '
    'decisions are interpretable, the aggregate prediction from hundreds of sequential trees defies human '
    'comprehension, and SHAP-based post-hoc explanations provide only approximate feature attributions '
    'rather than exact decision logic.'
)

doc.add_heading('2.2.2  Support Vector Machines', level=3)

add_para(
    'Support Vector Machines (SVMs) find optimal hyperplanes that maximize the margin between classes '
    'in a transformed feature space. Kernel functions (RBF, polynomial) enable SVMs to capture non-linear '
    'decision boundaries. Mukkamala et al. (2002) applied SVMs to the KDD Cup 99 dataset and reported '
    'detection rates exceeding 99% for DoS attacks, though performance on rarer attack types was '
    'significantly lower. SVMs scale poorly to large datasets due to their O(n^2) to O(n^3) training '
    'complexity, limiting their applicability to real-time IDS where training data accumulates continuously. '
    'Additionally, SVMs with non-linear kernels are opaque: the decision boundary exists in a transformed '
    'high-dimensional space that cannot be directly inspected or interpreted.'
)

doc.add_heading('2.2.3  Bayesian and Probabilistic Methods', level=3)

add_para(
    'Naive Bayes classifiers have been applied to IDS as computationally efficient baseline models that '
    'provide probabilistic class predictions. Despite their strong independence assumption, Naive Bayes '
    'classifiers achieve reasonable performance on network traffic data where features such as packet rate, '
    'byte rate, and duration are somewhat conditionally independent given the attack type. Hidden Markov '
    'Models (HMMs) have been used to model temporal sequences of network events, capturing state transitions '
    'that correspond to multi-stage attack patterns. Bayesian networks offer a more flexible probabilistic '
    'framework that can encode domain knowledge about dependencies between network features and attack types.'
)

add_table(
    ['Method', 'Accuracy', 'Interpretability', 'Training Speed', 'Key Limitation'],
    [
        ['Decision Tree', '85-92%', 'High', 'Fast', 'Overfitting'],
        ['Random Forest', '92-97%', 'Medium', 'Medium', 'Partial opacity'],
        ['XGBoost/LightGBM', '95-99%', 'Low', 'Medium', 'Black-box ensemble'],
        ['SVM (RBF)', '90-96%', 'Low', 'Slow', 'Scalability'],
        ['Naive Bayes', '80-88%', 'High', 'Very Fast', 'Independence assumption'],
        ['MLP', '93-97%', 'Very Low', 'Slow', 'Complete opacity'],
    ]
)
add_figure_ref('Table 2.1: Summary of ML-Based IDS Approaches and Their Characteristics')

# 2.3
doc.add_heading('2.3  Deep Learning and Neural Network-Based IDS', level=2)

add_para(
    'Deep learning has transformed multiple fields of AI and has been extensively applied to intrusion '
    'detection with notable performance improvements. This section reviews the major deep learning '
    'architectures that have been applied to IDS and evaluates their suitability for operational deployment.'
)

doc.add_heading('2.3.1  Convolutional Neural Networks', level=3)

add_para(
    'Convolutional Neural Networks (CNNs), originally designed for image recognition, have been adapted '
    'for IDS by reshaping network flow features into grid-like structures. Wang et al. (2017) converted '
    'raw packet bytes into grayscale images and applied CNN classifiers, achieving high accuracy on the '
    'ISCX dataset. This approach leverages the CNN\'s ability to learn hierarchical spatial features, '
    'analogous to texture patterns in images, from the byte distributions in network traffic. Kim et al. '
    '(2020) proposed a 1D-CNN architecture that operates directly on sequential flow features, achieving '
    'F1 scores of 0.96 on CIC-IDS2017 with significantly reduced inference time compared to 2D-CNN '
    'approaches. While CNNs achieve impressive accuracy, their convolutional filters and learned feature '
    'maps are difficult to interpret in terms of network security semantics, and their computational '
    'requirements present challenges for real-time deployment at network line speeds.'
)

doc.add_heading('2.3.2  Recurrent Neural Networks and LSTMs', level=3)

add_para(
    'Recurrent Neural Networks (RNNs) and their gated variants, particularly Long Short-Term Memory '
    '(LSTM) networks, are naturally suited to sequential data and have been applied to IDS for modeling '
    'temporal dependencies in network traffic. Yin et al. (2017) demonstrated that LSTM-based IDS '
    'outperformed traditional ML methods on the NSL-KDD dataset, particularly for attack types with '
    'temporal patterns such as slow-rate DoS and multi-stage intrusions. Attention mechanisms, borrowed '
    'from natural language processing, have been integrated into LSTM-based IDS to focus the network\'s '
    'attention on the most relevant time steps within a flow sequence. However, the sequential nature '
    'of RNN processing introduces latency that is problematic for real-time detection at high traffic '
    'volumes, and the learned temporal representations remain opaque.'
)

doc.add_heading('2.3.3  Autoencoders and Generative Models', level=3)

add_para(
    'Autoencoders have been widely applied to anomaly-based IDS, where they learn compact representations '
    'of normal network traffic and detect anomalies as inputs with high reconstruction error. Variational '
    'Autoencoders (VAEs) provide a probabilistic framework for anomaly scoring, while Generative Adversarial '
    'Networks (GANs) have been used both for data augmentation to address class imbalance in training data '
    'and for detecting adversarially generated traffic. Zenati et al. (2018) proposed a BiGAN-based anomaly '
    'detection approach that achieved competitive performance with supervised methods without requiring '
    'labeled attack data. The unsupervised nature of autoencoder-based detection aligns with the practical '
    'reality that labeled attack data is scarce and expensive to produce, but the learned latent '
    'representations and reconstruction thresholds lack the interpretability needed for forensic analysis.'
)

doc.add_heading('2.3.4  Transformer Architectures', level=3)

add_para(
    'More recently, transformer architectures have been explored for network traffic classification. '
    'The self-attention mechanism enables transformers to capture long-range dependencies between flow '
    'features without the sequential processing bottleneck of RNNs. Wu et al. (2022) proposed a '
    'flow-level transformer that processes sequences of network flows and achieved state-of-the-art '
    'performance on the CSE-CIC-IDS-2018 dataset. However, transformers require large amounts of '
    'training data to generalize effectively and have high computational costs during both training '
    'and inference. The attention weights provide some interpretability by indicating which flows or '
    'features influenced the classification, but this attention-based interpretability has been shown '
    'to be unreliable as a faithful explanation of model behavior.'
)

# 2.4
doc.add_heading('2.4  Reinforcement Learning in Cybersecurity', level=2)

add_para(
    'Reinforcement learning (RL) offers a fundamentally different approach to intrusion detection '
    'compared to supervised learning. Rather than learning a static mapping from features to labels, '
    'RL agents learn sequential decision policies through interaction with an environment, receiving '
    'reward signals that guide policy improvement. This section reviews the application of RL to '
    'cybersecurity generally and intrusion detection specifically.'
)

doc.add_heading('2.4.1  Foundations of Reinforcement Learning', level=3)

add_para(
    'Reinforcement learning formalizes sequential decision-making as a Markov Decision Process (MDP) '
    'defined by a tuple (S, A, T, R, gamma), where S is the state space, A is the action space, T is '
    'the transition function, R is the reward function, and gamma is the discount factor. The agent\'s '
    'objective is to learn a policy pi(a|s) that maximizes the expected cumulative discounted reward '
    'E[sum(gamma^t * R_t)]. Value-based methods such as Q-learning (Watkins & Dayan, 1992) estimate '
    'the action-value function Q(s, a) that represents the expected return of taking action a in state '
    's and following the optimal policy thereafter. Policy gradient methods such as REINFORCE (Williams, '
    '1992) and actor-critic methods such as PPO (Schulman et al., 2017) directly optimize the policy '
    'parameters to maximize expected return.'
)

add_para(
    'The exploration-exploitation trade-off is central to RL: the agent must balance exploiting its '
    'current knowledge to take high-reward actions with exploring new actions that might yield better '
    'long-term outcomes. Epsilon-greedy exploration, where the agent takes a random action with '
    'probability epsilon and the greedy action with probability 1 - epsilon, is the most common '
    'exploration strategy. The epsilon value is typically decayed over training, shifting from '
    'exploration to exploitation as the agent\'s policy improves.'
)

doc.add_heading('2.4.2  RL for Network Security', level=3)

add_para(
    'Servin and Kudenko (2008) were among the first to apply RL to intrusion detection, using '
    'multi-agent Q-learning where individual agents monitored different network segments and shared '
    'information to detect distributed attacks. Their work demonstrated that RL agents could learn '
    'effective detection policies through reward signals derived from correct and incorrect classifications. '
    'Subsequent research has explored various RL formulations for network security tasks.'
)

add_para(
    'Xu and Luo (2018) proposed a DQN-based IDS that frames intrusion detection as a sequential '
    'decision problem where the agent observes network flow features and decides whether to alert. '
    'The deep Q-network enables the agent to generalize across the high-dimensional feature space '
    'without manual feature engineering. Lopez-Martin et al. (2019) applied actor-critic methods to '
    'IDS and demonstrated that policy gradient approaches could achieve faster convergence than '
    'value-based methods on the NSL-KDD dataset. Caminero et al. (2019) explored adversarial RL '
    'for IDS, where an attacker agent learns to evade detection while the defender agent learns to '
    'detect evasive attacks, resulting in more robust detection policies.'
)

add_para(
    'Beyond intrusion detection, RL has been applied to automated penetration testing (Ghanem & '
    'Chen, 2020), network access control (Hu et al., 2020), and adaptive firewall rule management '
    '(Nguyen & Reddi, 2019). These applications share the characteristic of requiring sequential '
    'decisions in dynamic environments where the optimal policy changes as the threat landscape evolves. '
    'The common thread across these applications is the tension between the adaptability of RL and the '
    'safety requirements of security-critical systems.'
)

add_table(
    ['Study', 'RL Algorithm', 'Application', 'Dataset', 'Key Finding'],
    [
        ['Servin & Kudenko (2008)', 'Multi-agent Q-learning', 'Distributed IDS', 'Custom', 'Agents learn coordinated detection'],
        ['Xu & Luo (2018)', 'DQN', 'Flow classification', 'NSL-KDD', 'Generalizes across features'],
        ['Lopez-Martin et al. (2019)', 'Actor-Critic', 'Alert generation', 'NSL-KDD', 'Faster convergence than DQN'],
        ['Caminero et al. (2019)', 'Adversarial RL', 'Evasion-robust IDS', 'UNSW-NB15', 'Robust to adversarial attacks'],
        ['ASRRL (This work)', 'Shielded Q-learning', 'Safe classification', 'Multiple', 'Formal safety guarantees'],
    ]
)
add_figure_ref('Table 2.2: Reinforcement Learning Applications in Cybersecurity')

doc.add_heading('2.4.3  Safe Reinforcement Learning', level=3)

add_para(
    'Safe RL has emerged as a critical subfield addressing the deployment of RL agents in safety-critical '
    'domains. The core challenge is ensuring that the agent\'s policy satisfies safety constraints during '
    'both learning and deployment. Garcia and Fernandez (2015) provide a comprehensive taxonomy of safe RL '
    'approaches, categorizing them into constrained optimization methods, which incorporate safety constraints '
    'into the optimization objective; risk-sensitive methods, which modify the reward function to penalize '
    'risky actions; and shielding methods, which use external verification to filter unsafe actions.'
)

add_para(
    'Constrained Markov Decision Processes (CMDPs) extend the standard MDP framework by adding cost '
    'functions that the policy must keep below specified thresholds. Altman (1999) established the '
    'theoretical foundations of CMDPs, and recent work has developed practical algorithms for solving '
    'CMDPs, including constrained policy optimization (CPO) by Achiam et al. (2017) and Lagrangian-based '
    'methods. However, CMDP approaches provide probabilistic rather than deterministic safety guarantees: '
    'they ensure that expected constraint violations are below thresholds but do not guarantee that '
    'individual actions are safe.'
)

add_para(
    'Shielding approaches provide the strongest safety guarantees by interposing a verification layer '
    'between the RL agent and the environment. Alshiekh et al. (2018) introduced the concept of RL '
    'with formal guarantees through reactive synthesis, where a shield constructed from a formal '
    'specification monitors the agent\'s actions and substitutes safe alternatives when necessary. '
    'Jansen et al. (2020) extended this approach to partially observable settings. The ASRRL framework '
    'builds on the shielding paradigm by using Z3 satisfiability solving rather than reactive synthesis '
    'for shield construction, enabling dynamic constraint evolution through DBSCAN-based pattern '
    'discovery. This distinction is significant: while reactive synthesis shields are fixed at design '
    'time, Z3-based shields can incorporate new constraints as the system discovers novel attack patterns.'
)

# 2.5
doc.add_heading('2.5  Symbolic AI and Formal Verification', level=2)

add_para(
    'Symbolic AI, rooted in the logic-based paradigm of early artificial intelligence research, represents '
    'knowledge as explicit symbols and rules that are manipulated through formal inference procedures. '
    'While the dominance of statistical machine learning has relegated symbolic approaches to a secondary '
    'role in many application domains, the cybersecurity domain retains strong motivations for symbolic '
    'reasoning, including the need for verifiable security policies, auditable decision processes, and '
    'provable correctness guarantees.'
)

add_para(
    'Satisfiability Modulo Theories (SMT) solvers, of which Z3 (de Moura & Bjorner, 2008) is the most '
    'widely used, extend Boolean satisfiability (SAT) solving to first-order logic over theories including '
    'linear arithmetic, arrays, bit-vectors, and uninterpreted functions. SMT solvers have been extensively '
    'used in software verification, program analysis, and symbolic execution. Their application to machine '
    'learning verification is more recent: Katz et al. (2017) used SMT solving to verify properties of '
    'deep neural networks through the Reluplex algorithm, and subsequent work has developed specialized '
    'solvers for neural network verification (Marabou, alpha-beta-CROWN).'
)

add_para(
    'The application of formal verification to decision trees is particularly natural, as decision tree '
    'paths are inherently logical expressions: a path from root to leaf consists of a conjunction of '
    'threshold comparisons on feature values, and the set of all paths forms a disjunctive normal form '
    'representation of the classification function. This logical structure can be directly encoded as '
    'Z3 constraints, enabling efficient verification of properties such as: "for all inputs satisfying '
    'condition C, the tree predicts class L." The ASRRL framework exploits this correspondence by '
    'extracting Z3 constraints from decision tree paths and using them to verify RL actions.'
)

add_para(
    'Formal methods have a long history in cybersecurity, particularly in protocol verification and '
    'access control policy analysis. The use of model checking to verify network protocols (Clarke et '
    'al., 1999), theorem proving for cryptographic protocol analysis (Blanchet, 2001), and SAT solving '
    'for firewall rule conflict detection (Al-Shaer & Hamed, 2004) demonstrate the value of formal '
    'approaches in security-critical domains. The ASRRL framework extends this tradition by applying '
    'formal verification to the classification decisions of an ML-based IDS, bridging the gap between '
    'the adaptability of learning-based approaches and the rigor of formal methods.'
)

# 2.6
doc.add_heading('2.6  Neuro-Symbolic Integration', level=2)

add_para(
    'The integration of neural (statistical) and symbolic (logical) AI approaches has received growing '
    'attention as researchers seek to combine the learning capabilities of neural networks with the '
    'reasoning, interpretability, and verifiability of symbolic systems. Garcez et al. (2019) identify '
    'neuro-symbolic AI as a key pathway toward more robust and trustworthy AI systems, arguing that '
    'purely statistical approaches lack the compositional reasoning and systematic generalization '
    'capabilities that symbolic representations provide.'
)

add_para(
    'Several integration paradigms have been proposed. In the "neural for perception, symbolic for '
    'reasoning" paradigm, neural networks extract features from raw data while symbolic reasoners '
    'apply logical rules to the extracted features. DeepProbLog (Manhaeve et al., 2018) integrates '
    'neural networks with probabilistic logic programming, enabling end-to-end training of systems '
    'that combine perception and reasoning. In the "learning symbolic rules" paradigm, neural networks '
    'are used to learn symbolic rules from data, which can then be inspected and verified. Neural '
    'Logic Machines (Dong et al., 2019) learn first-order logic rules through differentiable '
    'programming, producing interpretable rule sets that generalize beyond the training distribution.'
)

add_para(
    'The ASRRL framework adopts a pragmatic integration approach: decision trees serve as the symbolic '
    'backbone, providing both classification predictions and a logical representation of the detection '
    'policy that can be extracted as Z3 constraints. The RL agent operates within the symbolic framework, '
    'using decision tree leaf IDs as its state representation and the Z3 constraint system as a safety '
    'shield. This design preserves the full interpretability of the decision tree while adding the '
    'adaptability of RL, and the Z3 verification ensures that the RL agent\'s learned policy remains '
    'consistent with the symbolic detection logic. This integration avoids the computational overhead '
    'of differentiable symbolic reasoning while providing concrete, verifiable safety guarantees.'
)

# 2.7
doc.add_heading('2.7  Adaptive and Self-Evolving Systems', level=2)

add_para(
    'The concept of adaptive security systems that evolve their detection capabilities in response to '
    'changing threat landscapes has been a longstanding goal in IDS research. Early adaptive systems '
    'focused on online learning algorithms that incrementally update model parameters as new data '
    'arrives. Eskin et al. (2002) proposed an adaptive anomaly detection framework using online '
    'clustering that updates cluster definitions as new normal behavior patterns emerge. Wang et al. '
    '(2006) developed an adaptive payload-based anomaly detector that automatically adjusts its '
    'detection thresholds based on traffic characteristics.'
)

add_para(
    'More recent approaches have explored ensemble-based adaptation, where multiple models trained on '
    'different time periods are combined, with weights adjusted based on each model\'s recent performance. '
    'Ditzler et al. (2015) provide a comprehensive survey of learning under concept drift, identifying '
    'strategies including sliding windows, decay functions, and drift detection tests. The challenge of '
    'adaptation in IDS is compounded by the adversarial nature of the environment: unlike natural concept '
    'drift in domains such as weather prediction, attack evolution is deliberately designed to exploit '
    'weaknesses in detection systems. This adversarial dimension necessitates adaptation mechanisms that '
    'are themselves robust to manipulation.'
)

add_para(
    'The ASRRL framework implements adaptation through three complementary mechanisms: RL policy updates '
    'that adjust the agent\'s detection strategy based on reward feedback, DBSCAN-based novel pattern '
    'detection that identifies emerging attack types from misclassified flows, and adaptive buffer '
    'management that adjusts the temporal granularity of traffic analysis based on current traffic '
    'characteristics. The combination of these mechanisms enables multi-scale adaptation: the RL policy '
    'adapts at the decision level, the constraint system adapts at the knowledge level, and the buffer '
    'adapts at the data processing level.'
)

# 2.8
doc.add_heading('2.8  Benchmark Datasets for IDS Research', level=2)

add_para(
    'The evaluation of IDS requires representative datasets that capture the complexity and diversity '
    'of real-world network traffic. This section reviews the major benchmark datasets used in IDS '
    'research and discusses their characteristics, strengths, and limitations.'
)

doc.add_heading('2.8.1  UNSW-NB15', level=3)

add_para(
    'The UNSW-NB15 dataset was created by Moustafa and Slay (2015) at the Australian Centre for Cyber '
    'Security using the IXIA PerfectStorm tool to generate a mixture of modern normal and attack '
    'activities. The dataset contains 2,540,044 records with 49 features and nine attack categories: '
    'Fuzzers, Analysis, Backdoors, DoS, Exploits, Generic, Reconnaissance, Shellcode, and Worms. The '
    'dataset includes both binary labels (normal/attack) and multi-class labels for attack type '
    'classification. Approximately 30% of records are attacks, providing a moderately imbalanced '
    'dataset. UNSW-NB15 is widely regarded as a more realistic alternative to the older KDD Cup 99 '
    'dataset, as it was generated using contemporary attack tools and protocols.'
)

doc.add_heading('2.8.2  CSE-CIC-IDS-2018', level=3)

add_para(
    'The CSE-CIC-IDS-2018 dataset was developed by the Canadian Institute for Cybersecurity in '
    'collaboration with the Communications Security Establishment. The dataset was generated from a '
    'controlled network environment simulating enterprise infrastructure with multiple victim machines '
    'and attack agents. It includes seven attack scenarios: Brute Force, Heartbleed, Botnet, DDoS, '
    'Web attacks (SQL injection, XSS), and infiltration. The dataset is characterized by a relatively '
    'low attack ratio (approximately 15%), reflecting the natural class imbalance in production '
    'environments. Features are extracted using CICFlowMeter and include 80 bidirectional flow features '
    'such as flow duration, packet length statistics, flag counts, and inter-arrival time distributions.'
)

doc.add_heading('2.8.3  CIC-IDS2017', level=3)

add_para(
    'The CIC-IDS2017 dataset (Sharafaldin et al., 2018) was generated over five days, with each day '
    'featuring different attack scenarios: Monday (benign only), Tuesday (brute force via SSH and FTP), '
    'Wednesday (DoS and Heartbleed), Thursday (web attacks and infiltration), and Friday (botnet, '
    'port scan, DDoS). The dataset contains approximately 2.8 million records with 78 features extracted '
    'by CICFlowMeter. The attack ratio is approximately 20%, and the temporal structure enables evaluation '
    'of time-based detection strategies. CIC-IDS2017 has been widely used in the IDS literature and '
    'provides a common benchmark for comparing detection methods.'
)

add_table(
    ['Dataset', 'Year', 'Records', 'Features', 'Attack %', 'Attack Types', 'Generator'],
    [
        ['UNSW-NB15', '2015', '2.54M', '49', '30%', '9 categories', 'IXIA PerfectStorm'],
        ['CSE-CIC-IDS-2018', '2018', '16M+', '80', '15%', '7 scenarios', 'CICFlowMeter'],
        ['CIC-IDS2017', '2017', '2.8M', '78', '20%', '5 day scenarios', 'CICFlowMeter'],
    ]
)
add_figure_ref('Table 2.3: Benchmark Dataset Characteristics')

add_para(
    'While these datasets provide valuable benchmarks for IDS research, several limitations have been '
    'identified. McHugh (2000) and Tavallaee et al. (2009) criticized the KDD Cup 99 dataset for '
    'redundant records, unrealistic traffic distributions, and outdated attack types, leading to the '
    'development of the NSL-KDD variant and subsequently the UNSW and CIC datasets. However, even '
    'contemporary datasets face challenges: synthetic traffic generation may not capture the full '
    'complexity of production network behavior, controlled network environments may not represent the '
    'diversity of real enterprise networks, and the labeled ground truth may contain errors. These '
    'limitations motivate the multi-dataset evaluation approach adopted in this dissertation, which '
    'tests the ASRRL framework across three datasets with different characteristics to assess '
    'generalization beyond any single dataset\'s idiosyncrasies.'
)

# 2.9
doc.add_heading('2.9  Interpretability and Explainable AI in Security', level=2)

add_para(
    'The need for interpretable and explainable AI (XAI) in cybersecurity has been increasingly '
    'recognized by both the research community and regulatory bodies. Lipton (2018) distinguishes '
    'between transparency, where the model\'s internal workings are directly understandable, and '
    'post-hoc interpretability, where explanations are generated after the model makes a prediction. '
    'In cybersecurity, transparency is preferable because security analysts need to understand not '
    'just which features were important for a classification but the precise decision logic that led '
    'to the classification, enabling them to validate the reasoning and identify potential blind spots.'
)

add_para(
    'Post-hoc explanation methods such as LIME (Ribeiro et al., 2016) and SHAP (Lundberg & Lee, 2017) '
    'generate local or global explanations for black-box models by training interpretable surrogate '
    'models or computing Shapley values. While these methods have been applied to IDS (e.g., Mane & '
    'Rao, 2021), they have significant limitations in the security domain. First, they provide '
    'approximate explanations that may not faithfully represent the model\'s actual decision process. '
    'Second, they are computationally expensive, adding latency to each classification. Third, they '
    'cannot provide safety guarantees about the model\'s behavior on unseen inputs. Rudin (2019) '
    'argues persuasively that in high-stakes domains, inherently interpretable models should be '
    'preferred over post-hoc explanations of black-box models.'
)

add_para(
    'The ASRRL framework aligns with Rudin\'s prescription by using inherently interpretable decision '
    'trees as the symbolic backbone and encoding their decision logic as verifiable Z3 constraints. '
    'Every classification decision can be traced to a specific path through the decision tree, expressed '
    'as a conjunction of human-readable threshold comparisons on network features. The Z3 constraint '
    'system adds a formal verification dimension that post-hoc methods cannot provide: the ability to '
    'prove that the system\'s behavior satisfies specified safety properties for all inputs within a '
    'defined region of the feature space.'
)

# 2.10
doc.add_heading('2.10  Identified Research Gaps', level=2)

add_para(
    'The literature review reveals several significant gaps at the intersection of IDS, reinforcement '
    'learning, and symbolic AI that collectively motivate the ASRRL framework:'
)

gaps = [
    ('Gap 1: No formal safety guarantees for RL-based IDS. ',
     'Existing RL-based IDS lack formal verification mechanisms to ensure that learned policies '
     'do not include unsafe actions. While safe RL research has explored constrained optimization '
     'and shielding approaches, these have not been applied to network intrusion detection with '
     'SMT-based verification.'),
    ('Gap 2: Limited neuro-symbolic integration in IDS. ',
     'Despite the natural fit between symbolic security policy representations and data-driven '
     'detection, few systems integrate formal symbolic reasoning with machine learning for '
     'intrusion detection. Existing neuro-symbolic approaches focus on NLP and computer vision '
     'rather than cybersecurity.'),
    ('Gap 3: Inadequate adaptation to novel attacks. ',
     'Current ML-based IDS rely on supervised retraining to adapt to new attack types, requiring '
     'labeled examples of novel attacks. No existing framework combines unsupervised pattern '
     'discovery with automatic constraint synthesis to extend detection capabilities without manual '
     'intervention.'),
    ('Gap 4: Incomplete evaluation methodologies. ',
     'Most IDS evaluations focus narrowly on classification accuracy metrics and fail to assess '
     'adversarial robustness, concept drift resilience, explanation fidelity, or statistical '
     'significance. This limits the ability to draw reliable conclusions about system effectiveness.'),
]

for title, desc in gaps:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(title)
    run.bold = True
    p.add_run(desc)

add_table(
    ['Research Gap', 'Addressed By', 'ASRRL Component'],
    [
        ['No safety guarantees for RL', 'Z3 constraint verification', 'Safety Shield'],
        ['Limited neuro-symbolic IDS', 'DT + Z3 + RL integration', 'Full Framework'],
        ['No adaptive constraint synthesis', 'DBSCAN + Z3 pipeline', 'Pattern Detector'],
        ['Incomplete evaluation', '10+ evaluation dimensions', 'Evaluation Framework'],
    ]
)
add_figure_ref('Table 2.4: Research Gap Analysis Matrix')

# 2.11
doc.add_heading('2.11  Theoretical Framework', level=2)

add_para(
    'The theoretical foundation of the ASRRL framework draws on three established theoretical traditions: '
    'the formal verification paradigm from computer science, which provides the mathematical basis for '
    'the Z3 constraint system; the Markov Decision Process formalism from operations research and AI, '
    'which underpins the reinforcement learning component; and the density-based clustering theory from '
    'unsupervised learning, which grounds the DBSCAN-based pattern discovery mechanism.'
)

add_para(
    'The integration of these theoretical traditions is guided by the principle of defense in depth: '
    'multiple independent detection and verification mechanisms operate in concert, such that the failure '
    'of any single mechanism does not compromise the system\'s overall safety. The decision tree provides '
    'a baseline classification that captures the primary attack signatures in the training data. The RL '
    'agent refines this classification through adaptive policy learning, and the Z3 shield ensures that '
    'the RL agent\'s decisions remain consistent with the symbolic model\'s knowledge. The DBSCAN pattern '
    'detector extends the symbolic model\'s knowledge by discovering novel patterns, closing the adaptation '
    'loop. This layered architecture provides both redundancy and complementarity: each component addresses '
    'a different type of uncertainty, and their combination yields a system that is more robust than any '
    'individual component.'
)

add_para(
    'Formally, the ASRRL framework operates as a Constrained MDP (CMDP) where the state space S is '
    'defined by decision tree leaf assignments, the action space A = {ALLOW, BLOCK, UNKNOWN}, the '
    'transition function T is determined by the network traffic stream, the reward function R encodes '
    'the detection objective with asymmetric penalties for false positives and false negatives, and the '
    'constraint set C is dynamically constructed from decision tree paths and DBSCAN-discovered patterns. '
    'The Z3 solver serves as the constraint evaluator, determining for each (state, action) pair whether '
    'the action is consistent with the current constraint set. This formulation ensures that the RL '
    'agent\'s policy converges to a solution that is both reward-maximizing and constraint-satisfying, '
    'providing the formal safety guarantees absent from unconstrained RL approaches.'
)

page_break()

# ═══════════════════════════════════════════════════════════════════
# CHAPTER 3: METHODOLOGY
# ═══════════════════════════════════════════════════════════════════
doc.add_heading('CHAPTER 3', level=1)
doc.add_heading('METHODOLOGY', level=1)

add_para(
    'This chapter presents the detailed methodology for the design, implementation, and evaluation of '
    'the ASRRL framework. The chapter begins with the research design philosophy, then describes each '
    'component of the framework architecture in detail, including the data collection and preprocessing '
    'pipeline, feature engineering, symbolic reasoning, reinforcement learning, safety shielding, novel '
    'pattern detection, and adaptive buffer management. The chapter concludes with a comprehensive '
    'description of the evaluation methodology spanning ten distinct evaluation dimensions.'
)

# 3.1
doc.add_heading('3.1  Research Design and Philosophy', level=2)

add_para(
    'This research adopts a design science methodology (Hevner et al., 2004) that emphasizes the '
    'creation and evaluation of IT artifacts to solve identified organizational problems. The ASRRL '
    'framework constitutes the primary artifact, designed to address the research gaps identified in '
    'Chapter 2. The design science approach is appropriate because the research objective is not merely '
    'to understand existing phenomena but to create a novel system that provides capabilities not '
    'available in existing solutions.'
)

add_para(
    'The research follows an iterative design-evaluate cycle: each component of the framework was '
    'designed, implemented, and evaluated independently before integration into the complete system. '
    'This component-wise development strategy enables ablation analysis, where the contribution of each '
    'component can be isolated by comparing the full framework against variants with specific components '
    'removed. The evaluation employs a mixed-methods approach combining quantitative performance metrics '
    '(accuracy, precision, recall, F1 score) with qualitative analysis of detection interpretability and '
    'constraint meaningfulness.'
)

add_para(
    'The research is conducted as a controlled experiment using established benchmark datasets that '
    'provide standardized conditions for comparing the ASRRL framework against baseline methods. '
    'Three datasets with distinct characteristics (UNSW-NB15, CSE-CIC-IDS-2018, CIC-IDS2017) are used '
    'to evaluate generalization, and multiple evaluation dimensions (statistical significance, '
    'adversarial robustness, concept drift resilience) are employed to assess robustness beyond standard '
    'accuracy metrics. The experimental design controls for random variation through multiple trials '
    'with different random seeds and reports results as mean plus-minus standard deviation with '
    'statistical significance tests.'
)

# 3.2
doc.add_heading('3.2  ASRRL Framework Architecture', level=2)

add_para(
    'The ASRRL framework consists of six interconnected components organized into a pipeline '
    'architecture that processes network flows from ingestion to classification. The architecture is '
    'designed for modularity, enabling independent evaluation and replacement of individual components, '
    'and for extensibility, supporting the integration of additional detection methods without modifying '
    'the core pipeline. The following describes the end-to-end processing flow:'
)

add_para(
    'Stage 1 - Flow Ingestion and Buffering: Raw network flows enter the system through the Adaptive '
    'Buffer, which aggregates flows into analysis windows. The buffer size is dynamically adjusted by '
    'the RL agent based on traffic characteristics, expanding during high-variability periods to capture '
    'sufficient context and contracting during stable periods to reduce latency. The initial buffer size '
    'is 20 flows, with a minimum of 10 and maximum of 200.',
    bold=False
)

add_para(
    'Stage 2 - Normalization: The Normalizer applies Z-score standardization to the continuous flow '
    'features (flow_duration, pkt_rate, byte_rate, entropy), transforming them to zero mean and unit '
    'variance. This normalization is essential for both the symbolic model\'s feature comparison and '
    'the DBSCAN clustering, which relies on Euclidean distances. The normalizer also computes aggregate '
    'buffer statistics (mean_entropy, byte_variance, mean_pkt_rate, mean_byte_rate) that serve as '
    'inputs to the RL agent\'s state representation.',
    bold=False
)

add_para(
    'Stage 3 - Symbolic Classification: The Decision Tree classifier (max_depth=6, min_samples_leaf=15) '
    'generates predictions and confidence scores for each analysis window. The tree is trained on the '
    'seven-feature representation (flow_duration, pkt_rate, byte_rate, entropy, port_cat, size_cat, '
    'protocol) and produces probability estimates for the attack and benign classes. The confidence '
    'score represents the proportion of attack predictions in the window.',
    bold=False
)

add_para(
    'Stage 4 - Z3 Constraint Extraction: Decision tree paths from root to each leaf node are extracted '
    'and encoded as Z3 constraints. Each path consists of a conjunction of threshold comparisons '
    '(e.g., entropy > 0.6 AND pkt_rate > 1800 implies BLOCK). The Z3 solver verifies the satisfiability '
    'of each constraint and caches the result using MD5 hashing of the state vector for efficiency.',
    bold=False
)

add_para(
    'Stage 5 - RL Decision with Safety Shielding: The Q-learning agent observes the current state '
    '(decision tree leaf ID), proposes an action (ALLOW, BLOCK, or UNKNOWN), and submits it to the Z3 '
    'safety shield for verification. If the proposed action satisfies all relevant constraints, it is '
    'executed. If it violates any constraint, the shield selects the highest-Q-value action among those '
    'that satisfy all constraints. If no safe action exists, the shield defaults to UNKNOWN and logs '
    'the shielding event.',
    bold=False
)

add_para(
    'Stage 6 - Novel Pattern Detection: Every three training epochs, the DBSCAN pattern detector '
    'analyzes the buffer of misclassified flows, identifies clusters representing potential novel '
    'attack patterns, and extracts cluster centroids as representative attack signatures. These '
    'centroids are converted into new Z3 constraints and added to the constraint system, enabling '
    'the safety shield to protect against newly discovered attack types.',
    bold=False
)

# 3.3
doc.add_heading('3.3  Data Collection and Preprocessing', level=2)

add_para(
    'The ASRRL framework is evaluated on three benchmark datasets that collectively represent a diverse '
    'range of network traffic conditions and attack types. Each dataset provides flow-level features '
    'extracted from captured network traffic, with binary labels indicating benign or attack flows. '
    'The preprocessing pipeline maps dataset-specific feature names to a standardized seven-feature '
    'schema used by the ASRRL framework.'
)

doc.add_heading('3.3.1  Dataset Descriptions', level=3)

add_para(
    'The UNSW-NB15 dataset, generated by the Australian Centre for Cyber Security using the IXIA '
    'PerfectStorm tool, provides 2.54 million network flow records with 49 features across nine attack '
    'categories. The dataset exhibits a 70:30 benign-to-attack ratio, making it moderately imbalanced. '
    'For this study, the following feature mappings are applied: dur maps to flow_duration, spkts and '
    'dpkts are combined to compute pkt_rate as (spkts + dpkts) / dur, sbytes and dbytes are combined '
    'to compute byte_rate as (sbytes + dbytes) / dur, and ct_state_ttl serves as a proxy for entropy. '
    'The sport and dsport fields are categorized into port_cat (well-known: 0, registered: 1-3, '
    'dynamic: 4-5), smeansz and dmeansz provide size_cat (small: 0, medium: 1, large: 2, jumbo: 3), '
    'and proto maps to protocol (TCP: 0, UDP: 1, other: 2).'
)

add_para(
    'The CSE-CIC-IDS-2018 dataset, developed by the Canadian Institute for Cybersecurity, contains '
    'over 16 million records across ten days of simulated enterprise network activity. The dataset '
    'features seven attack scenarios with a 15% attack ratio, reflecting the severe class imbalance '
    'typical of production environments. Feature mappings include Flow Duration to flow_duration, '
    'Fwd Packets/s and Bwd Packets/s combined for pkt_rate, Flow Bytes/s for byte_rate, and '
    'Flow IAT Std as an entropy proxy. Port, packet size, and protocol features are derived from '
    'Destination Port, Average Packet Size, and Protocol columns respectively.'
)

add_para(
    'The CIC-IDS2017 dataset contains approximately 2.8 million records generated over five days, '
    'with each day featuring different attack types. The 80:20 benign-to-attack ratio and temporal '
    'organization of attacks (brute force, DoS, web attacks, botnet, DDoS across different days) '
    'enables both aggregate and temporal evaluation. Feature mappings follow the same schema as '
    'CSE-CIC-IDS-2018, as both datasets use CICFlowMeter for feature extraction.'
)

doc.add_heading('3.3.2  Synthetic Faithful Data Generation', level=3)

add_para(
    'To ensure reproducible experimentation across varying dataset sizes, the framework includes a '
    'synthetic data generation module that produces flows with statistical distributions matching '
    'the original datasets. For each dataset, benign and attack flows are generated from Gaussian '
    'distributions whose parameters (mean and standard deviation) are estimated from the original '
    'dataset statistics. The following table presents the generation parameters for each dataset:'
)

add_table(
    ['Dataset', 'Feature', 'Benign (mu, sigma)', 'Attack (mu, sigma)'],
    [
        ['UNSW-NB15', 'flow_duration', '(1500, 400)', '(700, 300)'],
        ['UNSW-NB15', 'pkt_rate', '(600, 250)', '(1400, 350)'],
        ['UNSW-NB15', 'byte_rate', '(4e5, 1.5e5)', '(9e5, 2.5e5)'],
        ['UNSW-NB15', 'entropy', '(0.40, 0.15)', '(0.75, 0.18)'],
        ['CSE-CIC-2018', 'flow_duration', '(3000, 800)', '(1200, 500)'],
        ['CSE-CIC-2018', 'pkt_rate', '(1200, 600)', '(2400, 700)'],
        ['CSE-CIC-2018', 'byte_rate', '(1e6, 7e5)', '(2.8e6, 1.2e6)'],
        ['CSE-CIC-2018', 'entropy', '(0.45, 0.20)', '(0.95, 0.25)'],
        ['CIC-IDS2017', 'flow_duration', '(2000, 700)', '(900, 500)'],
        ['CIC-IDS2017', 'pkt_rate', '(900, 400)', '(2000, 600)'],
        ['CIC-IDS2017', 'byte_rate', '(7e5, 3e5)', '(1.8e6, 7e5)'],
        ['CIC-IDS2017', 'entropy', '(0.42, 0.18)', '(0.85, 0.22)'],
    ]
)
add_figure_ref('Table 3.1: Synthetic Data Generation Parameters by Dataset')

add_para(
    'Categorical features (port_cat, size_cat, protocol) are generated using uniform random sampling '
    'from their respective ranges: port_cat in {0, 1, 2, 3, 4, 5}, size_cat in {0, 1, 2, 3}, and '
    'protocol in {0, 1, 2}. All continuous features are clipped to ensure non-negative values with a '
    'minimum of 0.05 for entropy. The synthetic generation preserves the class imbalance ratios of the '
    'original datasets: 30% attack for UNSW-NB15, 15% for CSE-CIC-IDS-2018, and 20% for CIC-IDS2017.'
)

# 3.4
doc.add_heading('3.4  Feature Engineering Pipeline', level=2)

add_para(
    'The ASRRL framework operates on a standardized seven-feature representation designed to capture '
    'the essential characteristics of network flows while remaining computationally efficient for '
    'real-time processing. The feature set comprises four continuous features that characterize the '
    'temporal and volumetric properties of flows and three categorical features that encode protocol '
    'and service information.'
)

add_table(
    ['Feature', 'Type', 'Range', 'Description', 'Security Relevance'],
    [
        ['flow_duration', 'Continuous', '[0, inf) ms', 'Total flow duration', 'Short flows indicate scans; long flows indicate tunneling'],
        ['pkt_rate', 'Continuous', '[0, inf) pkt/s', 'Packet transmission rate', 'High rates indicate flooding attacks'],
        ['byte_rate', 'Continuous', '[0, inf) B/s', 'Byte transmission rate', 'Anomalous rates indicate exfiltration'],
        ['entropy', 'Continuous', '[0.05, inf)', 'Information entropy', 'High entropy indicates encrypted/obfuscated payloads'],
        ['port_cat', 'Categorical', '{0-5}', 'Port number category', 'Unusual ports indicate backdoors'],
        ['size_cat', 'Categorical', '{0-3}', 'Packet size category', 'Size anomalies indicate tunneling'],
        ['protocol', 'Categorical', '{0-2}', 'Network protocol', 'Protocol misuse indicates evasion'],
    ]
)
add_figure_ref('Table 3.2: Feature Engineering Specifications')

add_para(
    'The normalization stage applies Z-score standardization (x_normalized = (x - mu) / sigma) to the '
    'four continuous features using scikit-learn\'s StandardScaler. The scaler is fit on the training '
    'partition and applied to both training and test partitions to prevent data leakage. The normalized '
    'features (z_flow_duration, z_pkt_rate, z_byte_rate, z_entropy) are used as inputs to the decision '
    'tree classifier and the DBSCAN clustering algorithm. Aggregate buffer statistics, computed from the '
    'normalized features within each buffer window, provide the RL agent\'s observation: mean_entropy '
    '(average z_entropy across flows in the buffer), byte_variance (variance of z_byte_rate), '
    'mean_pkt_rate (average z_pkt_rate), and mean_byte_rate (average z_byte_rate).'
)

add_para(
    'The feature engineering pipeline is intentionally compact, using seven features rather than the '
    '49-80 features available in the original datasets. This design choice is motivated by three '
    'considerations: (1) fewer features yield shallower, more interpretable decision trees with paths '
    'that can be meaningfully expressed as human-readable constraints; (2) Z3 constraint verification '
    'scales with the number of variables, and a compact feature set keeps verification within the 500ms '
    'timeout; (3) the selected features capture the fundamental volumetric, temporal, and protocol '
    'characteristics that distinguish benign from malicious flows, as confirmed by feature importance '
    'analysis across all three datasets. The framework\'s architecture supports straightforward extension '
    'to additional features by adding Z3 real variables and adjusting the feature name list.'
)

# 3.5
doc.add_heading('3.5  Symbolic Reasoning Component', level=2)

add_para(
    'The symbolic reasoning component of the ASRRL framework uses decision tree classifiers as the '
    'source of interpretable detection logic and the Z3 satisfiability modulo theories (SMT) solver '
    'as the formal verification engine. This section details the constraint extraction algorithm, the '
    'Z3 encoding, and the verification procedure.'
)

doc.add_heading('3.5.1  Decision Tree as Symbolic Model', level=3)

add_para(
    'The decision tree classifier serves as the symbolic backbone of the ASRRL framework. Trained with '
    'max_depth=6 and min_samples_leaf=15 using scikit-learn\'s DecisionTreeClassifier with random_state=42, '
    'the tree partitions the seven-dimensional feature space into leaf regions, each associated with a '
    'predicted class (benign or attack) and a class probability distribution. The depth limit of 6 '
    'balances model expressiveness with interpretability: each root-to-leaf path contains at most 6 '
    'threshold comparisons, producing decision rules that a security analyst can read and validate. '
    'The min_samples_leaf constraint of 15 prevents overfitting to individual flows, ensuring that each '
    'leaf represents a statistically meaningful cluster of traffic.'
)

add_para(
    'The tree is trained on 70% of the dataset with the remaining 30% reserved for evaluation. After '
    'training, the tree is used for two purposes: (1) generating classification predictions with '
    'probability estimates, and (2) providing the state representation for the RL agent through leaf ID '
    'assignment. Each flow is mapped to its corresponding decision tree leaf using the apply() method, '
    'and the leaf ID becomes the agent\'s state observation. This state representation is compact (the '
    'tree has at most 2^6 = 64 potential leaves, though pruning typically reduces this to 20-40 active '
    'leaves), interpretable (each leaf corresponds to a specific conjunction of feature conditions), and '
    'generalizable (flows in the same leaf share similar feature profiles regardless of their exact values).'
)

doc.add_heading('3.5.2  Z3 Constraint Extraction', level=3)

add_para(
    'The Z3 constraint extraction algorithm traverses the decision tree from root to each leaf, '
    'accumulating the split conditions along each path and encoding them as Z3 implications. The '
    'algorithm operates as follows:'
)

add_para(
    'For each leaf node L in the decision tree, let path(L) = {(f_1, t_1, d_1), (f_2, t_2, d_2), '
    '..., (f_k, t_k, d_k)} be the sequence of splits on the path from root to L, where f_i is the '
    'feature index, t_i is the threshold value, and d_i is the direction (left/right). The Z3 encoding '
    'is: Implies(And(c_1, c_2, ..., c_k), action == predicted_class(L)), where c_i = (feature[f_i] <= t_i) '
    'if d_i = left, or c_i = (feature[f_i] > t_i) if d_i = right.',
    italic=True
)

add_para(
    'The Z3 variables are declared as follows: seven Real variables for continuous and categorical '
    'features (flow_duration, pkt_rate, byte_rate, entropy, port_cat, size_cat, protocol) and one '
    'Int variable for the action code. The action is encoded as: ALLOW = 0, BLOCK = 1, UNKNOWN = 2. '
    'The constraint extraction is performed once per training epoch, after the decision tree is '
    'retrained on any newly available data. The extraction process also logs the total number of '
    'constraints at each epoch for monitoring the growth of the symbolic knowledge base.'
)

doc.add_heading('3.5.3  Constraint Verification', level=3)

add_para(
    'When the RL agent proposes an action for a given state, the Z3 solver verifies whether the action '
    'is consistent with the extracted constraints. The verification procedure assigns the current flow\'s '
    'feature values to the Z3 real variables, assigns the proposed action code to the action variable, '
    'and checks whether the resulting constraint system is satisfiable. If the solver returns SAT '
    '(satisfiable), the action is deemed consistent with the symbolic model\'s knowledge, and execution '
    'proceeds. If the solver returns UNSAT (unsatisfiable), the action violates a constraint, and the '
    'safety shield activates to select an alternative action.'
)

add_para(
    'To ensure real-time performance, the verification procedure incorporates two optimization '
    'strategies. First, a 500-millisecond timeout is set on the Z3 solver; if verification does not '
    'complete within this window, the action is conservatively treated as unverified and the shield '
    'defaults to the UNKNOWN action. Second, verification results are cached using MD5 hashing of the '
    'state vector concatenated with the action code. The cache hit rate typically exceeds 80% after '
    'the initial training epoch, as many flows map to the same decision tree leaves and the same '
    'actions are repeatedly proposed for similar states.'
)

# 3.6
doc.add_heading('3.6  Reinforcement Learning Component', level=2)

add_para(
    'The reinforcement learning component employs tabular Q-learning to learn an adaptive classification '
    'policy within the state space defined by the decision tree\'s leaf assignments. This section details '
    'the MDP formulation, the Q-learning algorithm, the reward function, and the exploration strategy.'
)

doc.add_heading('3.6.1  MDP Formulation', level=3)

add_para(
    'The intrusion detection task is formulated as a Markov Decision Process (MDP) with the following '
    'components: The state space S is defined by the set of decision tree leaf IDs. Each flow is '
    'mapped to its leaf using the trained decision tree\'s apply() method, producing a state '
    'observation s = (leaf_id,). The action space A = {ALLOW (0), BLOCK (1), UNKNOWN (2)} represents '
    'the three possible classification decisions. ALLOW permits the flow, BLOCK flags the flow as an '
    'attack, and UNKNOWN defers the classification for further analysis. The transition function T(s\'|s, a) '
    'is determined by the network traffic stream: after classifying the current flow, the next flow in '
    'the stream determines the next state. The reward function R(s, a, y) depends on the action taken '
    'and the true label y of the flow.'
)

doc.add_heading('3.6.2  Q-Learning Algorithm', level=3)

add_para(
    'The Q-learning update rule is: Q(s, a) <- Q(s, a) + alpha * [r + gamma * max_a\' Q(s\', a\') - Q(s, a)], '
    'where alpha = 0.15 is the learning rate, gamma = 0.95 is the discount factor, r is the immediate '
    'reward, and s\' is the next state. The Q-table is initialized to zeros and indexed by (leaf_id, action) '
    'pairs. The tabular representation is feasible because the decision tree\'s depth constraint limits the '
    'number of leaves, and the three-action space keeps the table compact.'
)

add_table(
    ['Parameter', 'Symbol', 'Value', 'Justification'],
    [
        ['Learning rate', 'alpha', '0.15', 'Moderate adaptation speed'],
        ['Discount factor', 'gamma', '0.95', 'Strong future reward consideration'],
        ['Initial epsilon', 'eps_0', '0.20', 'Moderate initial exploration'],
        ['Epsilon decay', 'eps_decay', '0.999', 'Gradual shift to exploitation'],
        ['Minimum epsilon', 'eps_min', '0.01', 'Maintain minimal exploration'],
        ['Number of actions', '|A|', '3', 'ALLOW, BLOCK, UNKNOWN'],
    ]
)
add_figure_ref('Table 3.3: Q-Learning Hyperparameters')

doc.add_heading('3.6.3  Reward Function', level=3)

add_para(
    'The reward function is designed with asymmetric penalties that reflect the differential costs of '
    'misclassification in cybersecurity. Missing an attack (false negative) is typically more dangerous '
    'than generating a false alarm (false positive), so the penalty for allowing an attack (-3.0) is '
    'three times the penalty for blocking benign traffic (-1.0). The reward function also includes a '
    'shielding bonus (+0.5) when the safety shield corrects an unsafe action, incentivizing the agent '
    'to learn from shield activations.'
)

add_table(
    ['True Label', 'Action', 'Base Reward', 'Shield Bonus', 'Total'],
    [
        ['Attack (1)', 'BLOCK', '+2.0', '+0.5', '+2.5'],
        ['Attack (1)', 'ALLOW', '-3.0', 'N/A', '-3.0'],
        ['Attack (1)', 'UNKNOWN', '+0.5', 'N/A', '+0.5'],
        ['Benign (0)', 'ALLOW', '+1.0', 'N/A', '+1.0'],
        ['Benign (0)', 'BLOCK', '-1.0', 'N/A', '-1.0'],
        ['Benign (0)', 'UNKNOWN', '0.0', 'N/A', '0.0'],
        ['Uncertain (2)', 'UNKNOWN', '+1.5', 'N/A', '+1.5'],
        ['Uncertain (2)', 'Other', '-1.5', 'N/A', '-1.5'],
    ]
)
add_figure_ref('Table 3.4: Reward Function Specification')

doc.add_heading('3.6.4  Exploration Strategy', level=3)

add_para(
    'The agent employs epsilon-greedy exploration with exponential decay. At each time step, with '
    'probability epsilon the agent selects a random action from the action space, and with probability '
    '1 - epsilon it selects the action with the highest Q-value for the current state. The epsilon '
    'value decays by a factor of 0.999 per step, starting from 0.20 and decaying to a minimum of 0.01. '
    'This schedule ensures substantial exploration in early training (when the Q-table is poorly '
    'estimated) and near-complete exploitation in later training (when the Q-table has converged to '
    'accurate estimates). The minimum epsilon of 0.01 maintains a small probability of exploration even '
    'after convergence, enabling the agent to discover improved policies if the traffic distribution '
    'changes.'
)

# 3.7
doc.add_heading('3.7  Safety Shielding Mechanism', level=2)

add_para(
    'The safety shielding mechanism is the defining innovation of the ASRRL framework, providing formal '
    'guarantees that the RL agent\'s actions are consistent with the symbolic detection model\'s knowledge. '
    'The shield operates as an intermediary between the RL agent and the environment, intercepting every '
    'proposed action and verifying it against the Z3 constraint system before allowing execution.'
)

doc.add_heading('3.7.1  Shield Operation', level=3)

add_para(
    'The shielding process follows a three-step procedure for each classification decision. In the first '
    'step (action proposal), the RL agent observes the current state s (decision tree leaf ID) and '
    'proposes an action a_proposed based on its Q-table and exploration strategy. In the second step '
    '(verification), the shield submits a_proposed to the Z3 solver along with the current flow\'s '
    'feature values. The solver checks whether the constraint system is satisfiable with the proposed '
    'action. In the third step (action selection), if verification succeeds (SAT), the proposed action '
    'is executed. If verification fails (UNSAT), the shield enumerates all actions in the action space, '
    'verifies each, and selects the verified action with the highest Q-value. If no action can be '
    'verified, the shield defaults to UNKNOWN and sets a shielded flag for logging.'
)

add_para(
    'The shield activation rate is a key diagnostic metric. A high shield activation rate in early '
    'training is expected, as the RL agent\'s policy is initially random and frequently proposes actions '
    'inconsistent with the symbolic model. As training progresses and the agent learns a policy aligned '
    'with the constraint system, the shield activation rate should decrease. A persistently high shield '
    'activation rate may indicate a conflict between the reward function and the constraint system, '
    'requiring adjustment of either the reward structure or the constraint extraction parameters.'
)

doc.add_heading('3.7.2  Dynamic Constraint Evolution', level=3)

add_para(
    'Unlike traditional safety shields that are constructed from fixed specifications, the ASRRL shield\'s '
    'constraint set evolves over time as the DBSCAN pattern detector discovers novel attack patterns. When '
    'a new pattern is detected, the corresponding constraint is added to the Z3 system, immediately '
    'expanding the shield\'s verification scope. This dynamic evolution enables the shield to protect '
    'against attack types that were not present in the initial training data, providing a crucial '
    'capability for defending against zero-day attacks.'
)

add_para(
    'The constraint evolution is monotonic: new constraints are added but existing constraints are not '
    'removed. This design ensures that the shield\'s safety guarantees are cumulative: every attack pattern '
    'ever detected remains in the constraint system. The growth of the constraint set is logged at each '
    'epoch, enabling monitoring of the system\'s evolving knowledge base. In practice, the constraint count '
    'grows logarithmically with training duration, as novel patterns become increasingly rare as the system '
    'accumulates knowledge about the traffic environment.'
)

# 3.8
doc.add_heading('3.8  Novel Pattern Detection via DBSCAN', level=2)

add_para(
    'The DBSCAN (Density-Based Spatial Clustering of Applications with Noise) component provides the '
    'ASRRL framework with the ability to discover novel attack patterns without requiring labeled examples. '
    'DBSCAN was selected over other clustering algorithms (K-means, hierarchical clustering) for three '
    'reasons: (1) it does not require the number of clusters to be specified in advance, which is '
    'appropriate since the number of novel attack types is unknown; (2) it can identify clusters of '
    'arbitrary shape, accommodating the non-spherical distributions typical of attack traffic in '
    'normalized feature space; and (3) it naturally handles noise points (individual misclassifications '
    'that do not belong to any cluster), preventing spurious patterns from contaminating the constraint '
    'system.'
)

add_table(
    ['Parameter', 'Value', 'Justification'],
    [
        ['eps (neighborhood radius)', '1.5', 'Calibrated for StandardScaler-normalized features'],
        ['min_samples', '5', 'Minimum cluster size to ensure statistical significance'],
        ['buffer_size', '1000', 'Rolling window of recent misclassifications'],
        ['Detection frequency', 'Every 3 epochs', 'Balance between responsiveness and stability'],
        ['Feature space', '7-dimensional', 'All normalized flow features'],
    ]
)
add_figure_ref('Table 3.5: DBSCAN Configuration Parameters')

add_para(
    'The DBSCAN pattern detection pipeline operates as follows. During training, every flow that is '
    'misclassified (predicted label does not match true label) is added to a rolling buffer of size 1000. '
    'Every three epochs (starting from epoch 1), the buffer contents are clustered using DBSCAN with '
    'eps=1.5 and min_samples=5 on the StandardScaler-normalized feature space. For each cluster identified '
    '(excluding noise points labeled -1), the centroid is computed as the mean of all cluster members. '
    'Each centroid represents a previously undetected attack pattern. The centroid is converted into a '
    'Z3 constraint by creating a conjunction of conditions specifying that each feature is within one '
    'standard deviation of the centroid value, with the implied action set to BLOCK. This constraint is '
    'added to the Z3 system, and the n_novel counter is incremented.'
)

add_para(
    'The choice of eps=1.5 was determined empirically by evaluating cluster quality (silhouette score) '
    'across a range of values (0.5 to 3.0) on misclassified flows from the UNSW-NB15 dataset. The '
    'value of 1.5 provides a balance between sensitivity (detecting genuine clusters) and specificity '
    '(avoiding merging distinct patterns into a single cluster). The min_samples parameter of 5 ensures '
    'that at least 5 misclassified flows must share similar feature profiles before they are recognized '
    'as a pattern, preventing individual misclassifications from generating spurious constraints.'
)

# 3.9
doc.add_heading('3.9  Adaptive Buffer Management', level=2)

add_para(
    'The adaptive buffer is a dynamically sized sliding window that controls the temporal granularity '
    'of traffic analysis. Unlike fixed-window approaches that process a constant number of flows per '
    'analysis cycle, the adaptive buffer adjusts its size based on the current traffic characteristics, '
    'as directed by the RL agent and the pipeline\'s buffer management components.'
)

add_para(
    'The buffer is implemented as a collections.deque with a dynamic maxlen parameter. The initial size '
    'is 20 flows, configurable via the INIT_BUFFER parameter. The minimum size is 10 flows (MIN_BUFFER), '
    'ensuring sufficient data for meaningful statistical analysis. The maximum size is 200 flows '
    '(MAX_BUFFER), preventing excessive latency during high-traffic periods. The buffer supports three '
    'resize operations: increase (size += 10, capped at MAX_BUFFER), decrease (size -= 5, floored at '
    'MIN_BUFFER), and keep (no change). The asymmetric resize steps (increase by 10, decrease by 5) '
    'reflect a preference for larger buffers, which provide more statistical context at the cost of '
    'slightly increased latency.'
)

add_para(
    'The RL agents (Q-learning and PPO variants) determine the resize action based on the current '
    'buffer statistics. The Q-learning agent uses threshold-based heuristics: if mean_entropy > 1.1 '
    'or byte_variance > 5e10, the buffer is increased to capture more context during high-variability '
    'periods; if mean_entropy < 0.5 and byte_variance < 1e10, the buffer is decreased for faster '
    'response during stable periods. The PPO agent uses a combined scoring function: score = 0.6 * '
    'mean_entropy + 0.4 * (byte_variance / 5e10), with thresholds of 1.2 for increase and 0.6 for '
    'decrease. Both agents track resize events and history for post-hoc analysis of buffer adaptation '
    'behavior.'
)

add_para(
    'In the enhanced evaluation framework, the buffer management is further refined with dynamic '
    'threshold adaptation. The decision threshold (initially 0.5, range [0.3, 0.9]) adapts based on '
    'model confidence: if confidence exceeds 0.7, the threshold increases by 0.05 (becoming more '
    'conservative), and if confidence drops below 0.3, the threshold decreases by 0.05 (becoming more '
    'sensitive). This coupled buffer-threshold adaptation enables the system to adjust both the data '
    'aggregation granularity and the decision boundary simultaneously, providing multi-scale adaptation '
    'to changing traffic conditions. Temporal phase simulations demonstrate the system\'s ability to '
    'respond to attack bursts (70% attack rate) by expanding buffers and lowering thresholds, then '
    'recovering to normal operation as the attack subsides.'
)

# 3.10
doc.add_heading('3.10  Evaluation Methodology', level=2)

add_para(
    'The evaluation methodology encompasses ten distinct dimensions designed to provide a comprehensive '
    'assessment of the ASRRL framework\'s capabilities. This multi-dimensional approach addresses the '
    'limitation of single-metric evaluations prevalent in the IDS literature and establishes a rigorous '
    'standard for comparative analysis.'
)

doc.add_heading('3.10.1  Classification Performance Metrics', level=3)

add_para(
    'The primary classification metrics are computed from the confusion matrix of predictions versus '
    'true labels on the held-out test set (30% of data). Six metrics are reported:'
)

add_table(
    ['Metric', 'Formula', 'Interpretation'],
    [
        ['Accuracy', 'TP + TN / (TP + TN + FP + FN)', 'Overall correct classification rate'],
        ['Precision', 'TP / (TP + FP)', 'Fraction of attack alerts that are true attacks'],
        ['Recall', 'TP / (TP + FN)', 'Fraction of actual attacks that are detected'],
        ['F1 Score', '2 * Precision * Recall / (Precision + Recall)', 'Harmonic mean of precision and recall'],
        ['FPR', 'FP / (FP + TN)', 'Rate of benign traffic incorrectly flagged'],
        ['FNR', 'FN / (FN + TP)', 'Rate of attacks that go undetected'],
    ]
)
add_figure_ref('Table 3.6: Evaluation Metrics and Their Formulas')

doc.add_heading('3.10.2  Statistical Significance Testing', level=3)

add_para(
    'To ensure that observed performance differences are not attributable to random variation, the '
    'framework conducts 10 independent trials with different random seeds (seed = trial * 7 + 42). '
    'Results are reported as mean plus-minus standard deviation, and two non-parametric statistical tests '
    'are applied: the Wilcoxon signed-rank test for paired comparison of the ASRRL framework against '
    'each ablation variant, and the Mann-Whitney U test for independent comparison against baseline '
    'methods. A significance threshold of alpha = 0.05 is used, with Bonferroni correction for multiple '
    'comparisons where applicable.'
)

doc.add_heading('3.10.3  Component Ablation Analysis', level=3)

add_para(
    'The contribution of each framework component is assessed through ablation studies that compare the '
    'full ASRRL framework against four variants, each with one component removed: (1) No-Z3, which '
    'removes constraint extraction and verification, allowing the RL agent to operate without symbolic '
    'constraints; (2) No-DBSCAN, which removes novel pattern detection, preventing the system from '
    'discovering and constraining new attack types; (3) No-Shield, which removes the safety shielding '
    'mechanism, executing the RL agent\'s proposed actions without verification; and (4) No-RL, which '
    'uses only the decision tree\'s predictions without reinforcement learning refinement. These ablation '
    'variants isolate the marginal contribution of each component to overall system performance.'
)

doc.add_heading('3.10.4  Adversarial Robustness', level=3)

add_para(
    'Adversarial robustness is evaluated by perturbing test set features with Gaussian noise at seven '
    'intensity levels (epsilon in {0.0, 0.01, 0.05, 0.10, 0.20, 0.30, 0.50}). For each feature j and '
    'perturbation level epsilon, noise is sampled from N(0, epsilon * sigma_j) where sigma_j is the '
    'standard deviation of feature j in the training data. The perturbed test set is then classified by '
    'the trained ASRRL framework, and the F1 score is computed. This evaluation simulates adversarial '
    'attacks that manipulate traffic features to evade detection and quantifies the framework\'s '
    'resilience to such manipulation.'
)

doc.add_heading('3.10.5  Concept Drift Resilience', level=3)

add_para(
    'Concept drift is simulated by shifting the test set\'s attack distribution by drift * sigma in a '
    'random direction for each continuous feature, with additional Gaussian noise of magnitude '
    'drift * sigma * 0.5. Seven drift levels (0.0 to 1.0 in steps) are evaluated. The framework\'s F1 '
    'score at each drift level quantifies its ability to maintain detection performance as the attack '
    'distribution evolves. Additionally, temporal phase simulations with six distinct attack rate phases '
    '(normal at 5%, escalation at 25%, burst at 70%, recovery at 15%, second burst at 60%, and return '
    'to normal at 8%) evaluate the system\'s dynamic adaptation capabilities.'
)

doc.add_heading('3.10.6  Explanation Fidelity', level=3)

add_para(
    'Explanation fidelity measures the consistency between the Z3 constraint system\'s verification '
    'decisions and the final classification outputs. Three metrics are computed: fidelity (the proportion '
    'of flows where the Z3 decision agrees with the final classification, among flows where Z3 provides '
    'an opinion), coverage (the proportion of flows where at least one Z3 constraint applies), and opinion '
    'rate (the proportion of flows where the Z3 solver returns a definitive SAT or UNSAT result within '
    'the timeout). High fidelity indicates that the RL agent has learned a policy well-aligned with the '
    'symbolic model; high coverage indicates that the constraint system spans a large portion of the '
    'feature space; and high opinion rate indicates that the Z3 solver can efficiently verify most '
    'classification decisions.'
)

doc.add_heading('3.10.7  Scalability Analysis', level=3)

add_para(
    'Scalability is evaluated by training and testing the framework on datasets of five sizes: 1,000, '
    '5,000, 10,000, 25,000, and 50,000 samples. For each size, training time and throughput (flows per '
    'second) are measured. The same analysis is conducted for baseline models (Random Forest, XGBoost, '
    'LightGBM) to provide comparative scaling behavior. This evaluation identifies the practical '
    'throughput envelope within which the ASRRL framework can operate in real-time.'
)

doc.add_heading('3.10.8  Cross-Validation', level=3)

add_para(
    'Stratified 5-fold cross-validation is performed to assess the stability of performance estimates '
    'across different data partitions. Each fold maintains the class distribution of the full dataset. '
    'F1 scores are computed for each fold and reported as mean plus-minus standard deviation. Low variance '
    'across folds indicates that performance is not dependent on the specific train-test split, while '
    'high variance may indicate sensitivity to the data partition or model instability.'
)

add_table(
    ['Dimension', 'Method', 'Independent Variable', 'Dependent Variable'],
    [
        ['Performance', '70/30 split', 'Dataset', 'Accuracy, F1, FPR, FNR'],
        ['Significance', '10-trial evaluation', 'Random seed', 'Mean/Std, p-values'],
        ['Ablation', 'Component removal', 'Removed component', 'F1 score change'],
        ['Robustness', 'Feature perturbation', 'Epsilon (0-0.5)', 'F1 degradation'],
        ['Drift', 'Distribution shift', 'Drift level (0-1.0)', 'F1 under drift'],
        ['Fidelity', 'Z3 agreement analysis', 'Training epoch', 'Fidelity, coverage'],
        ['Scalability', 'Dataset size scaling', 'N (1K-50K)', 'Throughput (flows/s)'],
        ['Cross-validation', 'Stratified 5-fold', 'Fold index', 'F1 per fold'],
        ['Multi-class', 'Attack type labels', 'Attack category', 'Per-class F1'],
        ['Dynamic buffer', 'Phase simulation', 'Attack rate phase', 'Adaptation metrics'],
    ]
)
add_figure_ref('Table 3.7: Experimental Design Matrix')

# 3.11
doc.add_heading('3.11  Ethical Considerations', level=2)

add_para(
    'This research adheres to the ethical principles governing cybersecurity research and the responsible '
    'development of AI systems. Several ethical considerations are addressed in the design and execution '
    'of this study.'
)

add_para(
    'Data Privacy and Consent. The benchmark datasets used in this research (UNSW-NB15, CSE-CIC-IDS-2018, '
    'CIC-IDS2017) are publicly available datasets generated in controlled laboratory environments with '
    'synthetic traffic. No real user data or personally identifiable information is contained in these '
    'datasets. The synthetic data generation module used for scalability testing generates statistical '
    'replicas of the dataset distributions without reference to individual records, further ensuring that '
    'no privacy-sensitive information is used or generated during experimentation.'
)

add_para(
    'Dual-Use Considerations. Intrusion detection research inherently involves understanding attack '
    'methodologies, which could potentially be misused for offensive purposes. The ASRRL framework is '
    'designed exclusively for defensive application, specifically the detection and classification of '
    'network attacks. The framework does not generate attack traffic, exploit vulnerabilities, or provide '
    'offensive capabilities. The decision tree rules and Z3 constraints describe detection criteria rather '
    'than attack procedures. The publication of this research follows responsible disclosure practices, '
    'focusing on detection methodology without providing exploit details.'
)

add_para(
    'Algorithmic Fairness. While fairness is less commonly discussed in the cybersecurity context than in '
    'other AI application domains, the ASRRL framework\'s interpretable decision logic enables auditing '
    'for potential biases. The decision tree rules can be inspected to ensure that classification decisions '
    'are based on relevant network flow characteristics rather than potentially discriminatory attributes. '
    'The Z3 constraint system provides formal verification of decision properties, which could be extended '
    'to verify fairness constraints in addition to safety constraints.'
)

add_para(
    'Reproducibility. All experimental parameters, random seeds, dataset preparation procedures, and '
    'evaluation methodologies are documented in detail to enable independent reproduction of results. '
    'The source code implementing the ASRRL framework is made available in a public repository, and '
    'the synthetic data generation module enables researchers without access to the original datasets '
    'to reproduce experiments using statistically equivalent data.'
)

# ── REFERENCES (placeholder) ──
page_break()
doc.add_heading('REFERENCES', level=1)

refs = [
    'Achiam, J., Held, D., Tamar, A., & Abbeel, P. (2017). Constrained policy optimization. Proceedings of the 34th International Conference on Machine Learning, 22-31.',
    'Al-Shaer, E., & Hamed, H. (2004). Discovery of policy anomalies in distributed firewalls. IEEE INFOCOM 2004, 4, 2605-2616.',
    'Alshiekh, M., Bloem, R., Ehlers, R., Konighofer, B., Niekum, S., & Topcu, U. (2018). Safe reinforcement learning via shielding. AAAI Conference on Artificial Intelligence, 32(1).',
    'Altman, E. (1999). Constrained Markov decision processes. CRC Press.',
    'Amor, N. B., Benferhat, S., & Elouedi, Z. (2004). Naive Bayes vs decision trees in intrusion detection systems. ACM Symposium on Applied Computing, 420-424.',
    'Anderson, J. P. (1980). Computer security threat monitoring and surveillance. Technical Report, James P. Anderson Company.',
    'Blanchet, B. (2001). An efficient cryptographic protocol verifier based on Prolog rules. IEEE Computer Security Foundations Workshop, 82-96.',
    'Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5-32.',
    'Caminero, G., Lopez-Martin, M., & Carro, B. (2019). Adversarial environment reinforcement learning algorithm for intrusion detection. Computer Networks, 159, 96-109.',
    'Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 785-794.',
    'Clarke, E. M., Grumberg, O., & Peled, D. A. (1999). Model checking. MIT Press.',
    'de Moura, L., & Bjorner, N. (2008). Z3: An efficient SMT solver. International Conference on Tools and Algorithms for the Construction and Analysis of Systems, 337-340.',
    'Denning, D. E. (1987). An intrusion-detection model. IEEE Transactions on Software Engineering, SE-13(2), 222-232.',
    'Ditzler, G., Roveri, M., Alippi, C., & Polikar, R. (2015). Learning in nonstationary environments: A survey. IEEE Computational Intelligence Magazine, 10(4), 12-25.',
    'Dong, H., Mao, J., Lin, T., Wang, C., Li, L., & Zhou, D. (2019). Neural logic machines. International Conference on Learning Representations.',
    'Eskin, E., Arnold, A., Prerau, M., Portnoy, L., & Stolfo, S. (2002). A geometric framework for unsupervised anomaly detection. Applications of Data Mining in Computer Security, 77-101.',
    'Garcez, A. d., Gori, M., Lamb, L. C., Serafini, L., Spranger, M., & Tran, S. N. (2019). Neural-symbolic computing: An effective methodology for principled integration of machine learning and reasoning. Journal of Applied Logics, 6(4), 611-632.',
    'Garcia, J., & Fernandez, F. (2015). A comprehensive survey on safe reinforcement learning. Journal of Machine Learning Research, 16(1), 1437-1480.',
    'Ghanem, M. C., & Chen, T. M. (2020). Reinforcement learning for efficient network penetration testing. Information, 11(1), 6.',
    'Hevner, A. R., March, S. T., Park, J., & Ram, S. (2004). Design science in information systems research. MIS Quarterly, 28(1), 75-105.',
    'Hu, Z., Beuran, R., & Tan, Y. (2020). Automated penetration testing using deep reinforcement learning. IEEE European Symposium on Security and Privacy Workshops, 2-10.',
    'Jansen, N., Konighofer, B., Junges, S., Serban, A., & Bloem, R. (2020). Safe reinforcement learning using probabilistic shields. Concurrency and Computation: Practice and Experience, 32(17), e5824.',
    'Katz, G., Barrett, C., Dill, D. L., Julian, K., & Kochenderfer, M. J. (2017). Reluplex: An efficient SMT solver for verifying deep neural networks. International Conference on Computer Aided Verification, 97-117.',
    'Ke, G., Meng, Q., Finley, T., Wang, T., Chen, W., Ma, W., Ye, Q., & Liu, T. Y. (2017). LightGBM: A highly efficient gradient boosting decision tree. Advances in Neural Information Processing Systems, 30.',
    'Lashkari, A. H., Gil, G. D., Mamun, M. S. I., & Ghorbani, A. A. (2017). Characterization of tor traffic using time based features. International Conference on Information Systems Security and Privacy, 253-262.',
    'Lipton, Z. C. (2018). The mythos of model interpretability. Queue, 16(3), 31-57.',
    'Lopez-Martin, M., Carro, B., & Sanchez-Esguevillas, A. (2019). Application of deep reinforcement learning to intrusion detection for supervised problems. Expert Systems with Applications, 141, 112963.',
    'Lundberg, S. M., & Lee, S. I. (2017). A unified approach to interpreting model predictions. Advances in Neural Information Processing Systems, 30.',
    'Lunt, T. F., Tamaru, A., Gilham, F., Jagannathan, R., Neumann, P. G., Javitz, H. S., ... & Garvey, T. D. (1992). A real-time intrusion-detection expert system (IDES). Technical Report, SRI International.',
    'Manhaeve, R., Dumancic, S., Kimmig, A., Demeester, T., & De Raedt, L. (2018). DeepProbLog: Neural probabilistic logic programming. Advances in Neural Information Processing Systems, 31.',
    'Mane, S., & Rao, D. (2021). Explaining network intrusion detection system using explainable AI framework. arXiv preprint arXiv:2103.07110.',
    'McHugh, J. (2000). Testing intrusion detection systems: A critique of the 1998 and 1999 DARPA intrusion detection system evaluations. ACM Transactions on Information and System Security, 3(4), 262-294.',
    'Moustafa, N., & Slay, J. (2015). UNSW-NB15: A comprehensive data set for network intrusion detection systems. Military Communications and Information Systems Conference, 1-6.',
    'Mukkamala, S., Janoski, G., & Sung, A. (2002). Intrusion detection using neural networks and support vector machines. International Joint Conference on Neural Networks, 1702-1707.',
    'Nguyen, T. T., & Reddi, V. J. (2019). Deep reinforcement learning for cyber security. IEEE Transactions on Neural Networks and Learning Systems, 32(8), 3779-3795.',
    'Prokhorenkova, L., Gusev, G., Vorobev, A., Dorogush, A. V., & Gulin, A. (2018). CatBoost: Unbiased boosting with categorical features. Advances in Neural Information Processing Systems, 31.',
    'Ptacek, T. H., & Newsham, T. N. (1998). Insertion, evasion, and denial of service: Eluding network intrusion detection. Technical Report, Secure Networks.',
    'Quinlan, J. R. (1993). C4.5: Programs for machine learning. Morgan Kaufmann Publishers.',
    'Ribeiro, M. T., Singh, S., & Guestrin, C. (2016). "Why should I trust you?": Explaining the predictions of any classifier. ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 1135-1144.',
    'Roesch, M. (1999). Snort: Lightweight intrusion detection for networks. LISA, 99(1), 229-238.',
    'Rudin, C. (2019). Stop explaining black box machine learning models for high stakes decisions and use interpretable models instead. Nature Machine Intelligence, 1(5), 206-215.',
    'Schulman, J., Wolski, F., Dhariwal, P., Radford, A., & Klimov, O. (2017). Proximal policy optimization algorithms. arXiv preprint arXiv:1707.06347.',
    'Servin, A., & Kudenko, D. (2008). Multi-agent reinforcement learning for intrusion detection. Adaptive Agents and Multi-Agent Systems, 211-223.',
    'Sharafaldin, I., Lashkari, A. H., & Ghorbani, A. A. (2018). Toward generating a new intrusion detection dataset and intrusion traffic characterization. International Conference on Information Systems Security and Privacy, 108-116.',
    'Tavallaee, M., Bagheri, E., Lu, W., & Ghorbani, A. A. (2009). A detailed analysis of the KDD CUP 99 data set. IEEE Symposium on Computational Intelligence for Security and Defense Applications, 1-6.',
    'Watkins, C. J. C. H., & Dayan, P. (1992). Q-learning. Machine Learning, 8(3-4), 279-292.',
    'Williams, R. J. (1992). Simple statistical gradient-following algorithms for connectionist reinforcement learning. Machine Learning, 8(3-4), 229-256.',
    'Xu, X., & Luo, J. (2018). Deep reinforcement learning based network intrusion detection. IEEE International Conference on Software Quality, Reliability and Security, 538-542.',
    'Yin, C., Zhu, Y., Fei, J., & He, X. (2017). A deep learning approach for intrusion detection using recurrent neural networks. IEEE Access, 5, 21954-21961.',
    'Zenati, H., Foo, C. S., Lecouat, B., Manek, G., & Chandrasekhar, V. R. (2018). Efficient GAN-based anomaly detection. arXiv preprint arXiv:1802.06222.',
    'Zhang, J., & Zulkernine, M. (2006). A hybrid network intrusion detection technique using random forests. First International Conference on Availability, Reliability and Security, 262-269.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(ref)
    run.font.size = Pt(11)

# ── Save ──
out_path = '/home/user/adaptive-ids-asrrl/dissertation_chapters_1_3.docx'
doc.save(out_path)
print(f'Dissertation saved to {out_path}')
